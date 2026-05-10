"""KOSGEB AR-GE / Yenilik programme module.

Per docs/07 §6. Türkiye'nin ikinci en büyük KOBİ hibe kaynağı —
Türkçe-zorunlu, KOBİ-only, min 1 yıl şirket yaşı, 12-36 ay süre.
TÜBİTAK 1501'e benzer bir form yapısı ama daha kısa, daha az teknik
derinlik bekleniyor.

Section layout (KOSGEB KBS form):
  excellence: K1 Proje Konusu / K2 Yenilik Niteliği / K3 Uygulanacak Yöntem
  impact:     L1 Beklenen Ekonomik Kazanım / L2 Pazar Analizi
  implementation: M1 İş-Zaman Planı / M2 Bütçe / M3 Riskler

Budget structure: 4 kalem (Personel, Makine-Teçhizat-Yazılım, Hizmet
Alımı, Diğer) — TÜBİTAK'ın 6 kaleminden farklı, bu yüzden ayrı export
kuralları.

K2 (Yenilik Niteliği) en kritik bölüm — KOSGEB değerlendirmesinde en
ağırlıklı. Min 600 kelime şartı koyduk (TÜBİTAK 1501 B2'nin 800 kelimeden
biraz altı; KOSGEB değerlendiricileri TÜBİTAK kadar uzun beklemiyor).
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any, ClassVar

from docx import Document

from src.exports.md_to_docx import render_markdown
from src.programs.base import (
    BaseProgramModule,
    BriefField,
    BriefSchema,
    BriefSection,
    CallMetadata,
    ProgrammeLanguage,
    ValidationIssue,
)

_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / "templates" / "kosgeb_basvuru_2026.docx"
)
_PROMPTS_DIR = (
    Path(__file__).resolve().parent.parent.parent
    / "agents"
    / "prompts"
    / "kosgeb_arge"
)


_SUBSECTION_MAP: dict[str, list[str]] = {
    "excellence": [
        "K1_proje_konusu",
        "K2_yenilik_niteligi",
        "K3_uygulanacak_yontem",
    ],
    "impact": [
        "L1_beklenen_ekonomik_kazanim",
        "L2_pazar_analizi",
    ],
    "implementation": [
        "M1_is_zaman_plani",
        "M2_butce",
        "M3_riskler",
    ],
}

_SECTION_HEADERS: list[tuple[str, str]] = [
    ("excellence_md", "K. Proje Tanımı"),
    ("impact_md", "L. Etki ve Pazar"),
    ("implementation_md", "M. Uygulama"),
]

_BUDGET_CATEGORIES = ["Personel", "Makine-Teçhizat-Yazılım", "Hizmet Alımı", "Diğer"]

_TEMPLATE_LEADING_PARAGRAPHS_KEPT = 3


class KOSGEBARGEModule(BaseProgramModule):
    program_id: ClassVar[str] = "kosgeb_arge"
    name_tr: ClassVar[str] = "KOSGEB AR-GE / Yenilik Destek Programı"
    name_en: ClassVar[str] = "KOSGEB R&D / Innovation Support Programme"
    funder: ClassVar[str] = "KOSGEB"
    language: ClassVar[ProgrammeLanguage] = "tr"

    sections: ClassVar[list[str]] = ["excellence", "impact", "implementation"]
    subsection_map: ClassVar[dict[str, list[str]]] = _SUBSECTION_MAP

    duration_min_months: ClassVar[int] = 12
    duration_max_months: ClassVar[int] = 36
    requires_kobi: ClassVar[bool] = True
    min_company_age_years: ClassVar[int] = 1
    k2_min_words: ClassVar[int] = 600

    # ── Plugin API ─────────────────────────────────────────────────────

    def parse_call(self, call_text: str, call_metadata: dict[str, Any]) -> CallMetadata:
        del call_text
        return CallMetadata.model_validate({**call_metadata, "language": "tr"})

    def get_brief_schema(self) -> BriefSchema:
        return BriefSchema(
            sections=[
                BriefSection(
                    title_tr="Proje Özü",
                    title_en="Project Core",
                    fields=[
                        BriefField(
                            key="title",
                            label_tr="Proje başlığı",
                            label_en="Project title",
                            type="text",
                            max_length=200,
                        ),
                        BriefField(
                            key="problem_statement",
                            label_tr="Çözmeye çalıştığınız problem",
                            label_en="Problem you're solving",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="proposed_solution",
                            label_tr="Önerdiğiniz çözüm — yenilik niteliği",
                            label_en="Your proposed solution — innovation aspect",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="duration_months",
                            label_tr="Proje süresi (ay)",
                            label_en="Project duration (months)",
                            type="number",
                        ),
                    ],
                ),
                BriefSection(
                    title_tr="Şirket Bilgileri",
                    title_en="Company Information",
                    fields=[
                        BriefField(
                            key="company_age_years",
                            label_tr="Şirket yaşı (yıl)",
                            label_en="Company age (years)",
                            type="number",
                            help_text_tr="KOSGEB AR-GE için min 1 yıl gerekli.",
                            help_text_en="KOSGEB requires the company to be ≥1 year old.",
                        ),
                        BriefField(
                            key="is_sme",
                            label_tr="KOBİ mi?",
                            label_en="Is the company an SME?",
                            type="select",
                            options=[
                                {"value": True, "label_tr": "Evet", "label_en": "Yes"},
                                {"value": False, "label_tr": "Hayır", "label_en": "No"},
                            ],
                        ),
                    ],
                ),
            ]
        )

    def get_template_path(self) -> str:
        return str(_TEMPLATE_PATH)

    def get_prompt_path(self, agent_id: str) -> str:
        return str(_PROMPTS_DIR / agent_id / "v1.md")

    def validate_draft(
        self, draft: dict[str, Any], metadata: CallMetadata
    ) -> list[ValidationIssue]:
        del metadata
        issues: list[ValidationIssue] = []

        # K2 minimum length — most critical section, parallels TÜBİTAK B2.
        excellence_md = str(draft.get("excellence_md") or "")
        k2_text = _extract_subsection_body(excellence_md, "K2")
        if len(k2_text.split()) < self.k2_min_words:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="excellence",
                    code="k2_too_short",
                    message_tr=(
                        f"K2 (Yenilik Niteliği) en az {self.k2_min_words} kelime olmalı. "
                        "KOSGEB değerlendirmesinde en kritik bölüm."
                    ),
                    message_en=(
                        f"K2 (Innovation Quality) must be at least {self.k2_min_words} words. "
                        "Most critical section in KOSGEB eval."
                    ),
                    suggestion=(
                        "Mevcut çözümlerle karşılaştırmalı yenilik niteliğini "
                        "somut metriklerle anlatın."
                    ),
                )
            )

        brief = draft.get("brief") or {}

        # Duration check
        duration = int(brief.get("duration_months") or 0)
        if duration and not (
            self.duration_min_months <= duration <= self.duration_max_months
        ):
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="implementation",
                    code="invalid_duration",
                    message_tr=(
                        f"Proje süresi {duration} ay. {self.name_tr}'de "
                        f"{self.duration_min_months}-{self.duration_max_months} "
                        "ay arası zorunlu."
                    ),
                    message_en=(
                        f"Project duration {duration} months. {self.name_en} "
                        f"requires {self.duration_min_months}-"
                        f"{self.duration_max_months} months."
                    ),
                )
            )

        # KOBİ check (KOSGEB is SME-only by definition)
        if brief.get("is_sme") is False:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="implementation",
                    code="not_kobi",
                    message_tr=(
                        "KOSGEB AR-GE yalnızca KOBİ'lere açık. "
                        "Şirket KOBİ tanımına uymuyor."
                    ),
                    message_en=(
                        "KOSGEB R&D is open only to SMEs. "
                        "The company is not classified as an SME."
                    ),
                )
            )

        # Min company age
        try:
            company_age = float(brief.get("company_age_years") or 0)
        except (TypeError, ValueError):
            company_age = 0.0
        if 0 < company_age < self.min_company_age_years:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="implementation",
                    code="company_too_young",
                    message_tr=(
                        f"Şirket yaşı {company_age:.1f} yıl. "
                        f"KOSGEB AR-GE için min {self.min_company_age_years} yıl gerekli."
                    ),
                    message_en=(
                        f"Company is {company_age:.1f} years old. "
                        f"KOSGEB R&D requires ≥{self.min_company_age_years} years."
                    ),
                )
            )

        return issues

    def export_docx(self, proposal: dict[str, Any]) -> bytes:
        if not _TEMPLATE_PATH.is_file():
            raise FileNotFoundError(
                f"KOSGEB template missing: {_TEMPLATE_PATH}. Run "
                "`python -m src.programs.kosgeb_arge.template_builder` to regenerate."
            )

        doc = Document(str(_TEMPLATE_PATH))
        _strip_template_body(doc)
        _set_proposal_title(doc, str(proposal.get("title") or "İsimsiz Proje"))

        draft = proposal.get("draft") or {}
        for source_key, header in _SECTION_HEADERS:
            doc.add_heading(header, level=1)
            body = str(draft.get(source_key) or "").strip()
            if body:
                render_markdown(doc, body)
            else:
                doc.add_paragraph(f"[{source_key} bölümü henüz üretilmedi.]")

        # Budget table — KOSGEB has 4 categories vs TÜBİTAK's 6.
        budget = (proposal.get("budget") or {}).get("by_category") or {}
        _render_budget_table(doc, budget)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_xlsx_budget(self, proposal: dict[str, Any]) -> bytes | None:
        # KOSGEB KBS uses a portal form, not Excel. No XLSX export.
        del proposal
        return None


# ── Helpers ────────────────────────────────────────────────────────────


def _strip_template_body(doc: Any) -> None:
    """Remove placeholder paragraphs the template builder seeded."""

    body = doc.element.body
    children = list(body.iterchildren())
    paragraphs_seen = 0
    for child in children:
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            paragraphs_seen += 1
            if paragraphs_seen <= _TEMPLATE_LEADING_PARAGRAPHS_KEPT:
                continue
        if tag == "sectPr":
            continue
        body.remove(child)


def _set_proposal_title(doc: Any, title: str) -> None:
    for para in doc.paragraphs:
        if "[proje_basligi]" in para.text:
            for run in para.runs:
                if "[proje_basligi]" in run.text:
                    run.text = run.text.replace("[proje_basligi]", title)
            return


def _render_budget_table(
    doc: Any, by_category: dict[str, list[dict[str, Any]]]
) -> None:
    from docx.shared import Pt

    doc.add_heading("Bütçe Tablosu", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    for idx, label in enumerate(
        ["Gider Kalemi", "Açıklama", "Miktar", "Birim Fiyat (TL)", "Toplam (TL)"]
    ):
        headers[idx].text = label
        run = headers[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    grand_total = 0.0
    for category in _BUDGET_CATEGORIES:
        items = by_category.get(category) or []
        for item in items:
            row = table.add_row().cells
            row[0].text = category
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("quantity", ""))
            row[3].text = _fmt_money(item.get("unit_price"))
            total = float(item.get("total") or 0)
            row[4].text = _fmt_money(total)
            grand_total += total

    if grand_total > 0:
        total_row = table.add_row().cells
        total_row[0].text = "TOPLAM"
        total_row[0].paragraphs[0].runs[0].bold = True
        total_row[4].text = _fmt_money(grand_total)
        total_row[4].paragraphs[0].runs[0].bold = True


def _fmt_money(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return str(value)


def _extract_subsection_body(md: str, prefix: str) -> str:
    pattern = re.compile(rf"^##\s+{re.escape(prefix)}\b.*$", re.MULTILINE)
    match = pattern.search(md)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^##\s+", md[start:], re.MULTILINE)
    end = start + next_match.start() if next_match else len(md)
    return md[start:end].strip()


__all__ = ["KOSGEBARGEModule"]
