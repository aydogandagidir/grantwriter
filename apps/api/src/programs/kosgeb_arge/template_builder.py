"""Build the KOSGEB AR-GE / Yenilik DOCX scaffold programmatically.

Per docs/07 §6. Different section structure from TÜBİTAK (K/L/M
prefixes vs B/C/D), and a 4-category budget table vs TÜBİTAK's 6.
Run ``main()`` once to commit ``templates/kosgeb_basvuru_2026.docx``;
runtime export opens that committed binary.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentT
from docx.shared import Cm, Pt

PROGRAMME_NAME = "KOSGEB AR-GE / Yenilik Destek Programı"
SUBTITLE = "KBS — Başvuru Formu"

SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "K. Proje Tanımı",
        [
            ("K1", "Proje Konusu"),
            ("K2", "Yenilik Niteliği"),
            ("K3", "Uygulanacak Yöntem"),
        ],
    ),
    (
        "L. Etki ve Pazar",
        [
            ("L1", "Beklenen Ekonomik Kazanım"),
            ("L2", "Pazar Analizi"),
        ],
    ),
    (
        "M. Uygulama",
        [
            ("M1", "İş-Zaman Planı"),
            ("M2", "Bütçe"),
            ("M3", "Riskler"),
        ],
    ),
]

PLACEHOLDER_TEXT = "[Bu bölüm proje üretimi sırasında doldurulacaktır.]"


def build_template() -> DocumentT:
    """Return a fresh Document with the KOSGEB scaffold."""

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    title = doc.add_heading(PROGRAMME_NAME, level=0)
    for run in title.runs:
        run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run(SUBTITLE).italic = True

    title_p = doc.add_paragraph()
    title_p.add_run("Proje Başlığı: ").bold = True
    title_p.add_run("[proje_basligi]")

    for section_title, subsections in SECTIONS:
        doc.add_heading(section_title, level=1)
        for code, label in subsections:
            doc.add_heading(f"{code} {label}", level=2)
            doc.add_paragraph(PLACEHOLDER_TEXT)

    doc.add_heading("Bütçe Tablosu", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for idx, label in enumerate(
        ["Gider Kalemi", "Açıklama", "Miktar", "Birim Fiyat (TL)", "Toplam (TL)"]
    ):
        header[idx].text = label
        header[idx].paragraphs[0].runs[0].bold = True
    placeholder_row = table.add_row().cells
    placeholder_row[0].text = "[Personel|Makine-Teçhizat|Hizmet Alımı|Diğer]"
    placeholder_row[1].text = "[açıklama]"
    placeholder_row[2].text = "[adet]"
    placeholder_row[3].text = "[birim fiyat]"
    placeholder_row[4].text = "[toplam]"

    return doc


def to_bytes() -> bytes:
    buf = BytesIO()
    build_template().save(buf)
    return buf.getvalue()


def main(target: Path | None = None) -> Path:
    target = target or (
        Path(__file__).resolve().parent / "templates" / "kosgeb_basvuru_2026.docx"
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    build_template().save(str(target))
    return target


if __name__ == "__main__":
    written = main()
    print(f"wrote: {written}")


__all__ = ["PROGRAMME_NAME", "SECTIONS", "build_template", "main", "to_bytes"]
