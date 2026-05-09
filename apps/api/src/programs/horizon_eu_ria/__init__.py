"""Horizon Europe RIA / IA programme module.

Per docs/07 §3. Brief schema + section structure + validation rules
specific to HE Standard Application Form Part B. Lump Sum XLSX export
and DistinctivenessScorer integration land in S2.D7+ (export_xlsx_budget
returns ``None`` for now).
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from docx.document import Document as DocumentT
from docx.shared import Pt

from src.exports.md_to_docx import render_markdown
from src.programs.base import (
    BaseProgramModule,
    BriefField,
    BriefSchema,
    BriefSection,
    CallMetadata,
    ProgrammeLanguage,
    ValidationIssue,
)

_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "ria_part_b_2026.docx"
# Prompts are shared across HE RIA and (future) HE IA, hence the
# ``horizon_eu`` (not ``horizon_eu_ria``) directory.
_PROMPTS_DIR = Path(__file__).resolve().parent.parent.parent / "agents" / "prompts" / "horizon_eu"


_SUBSECTION_MAP: dict[str, list[str]] = {
    "excellence": [
        "1.1_objectives_and_ambition",
        "1.2_methodology",
        "1.3_state_of_the_art",
        "1.4_open_science",
    ],
    "impact": [
        "2.1_pathways_to_impact",
        "2.2_measures_to_maximise_impact",
        "2.3_summary_canvas",
    ],
    "implementation": [
        "3.1_work_plan_and_resources",
        "3.2_capacity_of_participants",
        "3.3_consortium_as_a_whole",
    ],
}

_SECTION_HEADERS: list[tuple[str, str]] = [
    ("excellence_md", "1. Excellence"),
    ("impact_md", "2. Impact"),
    ("implementation_md", "3. Implementation"),
]

_SUBSECTION_PREFIXES: dict[str, list[str]] = {
    "excellence": ["1.1", "1.2", "1.3", "1.4"],
    "impact": ["2.1", "2.2", "2.3"],
    "implementation": ["3.1", "3.2", "3.3"],
}


class HorizonEURIAModule(BaseProgramModule):
    program_id: ClassVar[str] = "horizon_eu_ria"
    name_tr: ClassVar[str] = "Horizon Europe RIA/IA"
    name_en: ClassVar[str] = "Horizon Europe RIA/IA"
    funder: ClassVar[str] = "European Commission"
    language: ClassVar[ProgrammeLanguage] = "en"

    sections: ClassVar[list[str]] = ["excellence", "impact", "implementation"]
    subsection_map: ClassVar[dict[str, list[str]]] = _SUBSECTION_MAP

    page_limit: ClassVar[int] = 45  # RIA / IA fixed limit
    min_partners: ClassVar[int] = 3
    min_countries: ClassVar[int] = 3
    words_per_page: ClassVar[int] = 500

    # ── Plugin API ─────────────────────────────────────────────────────

    def parse_call(self, call_text: str, call_metadata: dict[str, Any]) -> CallMetadata:
        del call_text  # CallAnalyst already parsed; we just normalise.
        merged = {**call_metadata}
        merged.setdefault("language", "en")
        merged.setdefault("page_limit", self.page_limit)
        return CallMetadata.model_validate(merged)

    def get_brief_schema(self) -> BriefSchema:
        # Verbatim from docs/07 §3.3 — the brief HE writers consume.
        trl_options: list[dict[str, Any]] = [
            {"value": i, "label": f"TRL {i}"} for i in range(1, 10)
        ]
        return BriefSchema(
            sections=[
                BriefSection(
                    title_tr="Proje Özü",
                    title_en="Project Core",
                    fields=[
                        BriefField(
                            key="title",
                            label_tr="Proje başlığı",
                            label_en="Project title",
                            type="text",
                            max_length=200,
                        ),
                        BriefField(
                            key="acronym",
                            label_tr="Akronim",
                            label_en="Acronym",
                            type="text",
                            max_length=15,
                            required=False,
                        ),
                        BriefField(
                            key="problem_statement",
                            label_tr="Çözmeye çalıştığınız problem (200-400 kelime)",
                            label_en="Problem you're solving (200-400 words)",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="proposed_solution",
                            label_tr="Önerdiğiniz çözüm — teknik yaklaşım",
                            label_en="Your proposed solution — technical approach",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="trl_current",
                            label_tr="Şu anki TRL",
                            label_en="Current TRL",
                            type="select",
                            options=trl_options,
                        ),
                        BriefField(
                            key="trl_target",
                            label_tr="Proje sonu hedef TRL",
                            label_en="Target TRL at project end",
                            type="select",
                            options=trl_options,
                        ),
                    ],
                ),
                BriefSection(
                    title_tr="Konsorsiyum",
                    title_en="Consortium",
                    fields=[
                        BriefField(
                            key="role",
                            label_tr="Sizin rolünüz",
                            label_en="Your role",
                            type="select",
                            options=[
                                {
                                    "value": "coordinator",
                                    "label_tr": "Koordinatör",
                                    "label_en": "Coordinator",
                                },
                                {
                                    "value": "partner",
                                    "label_tr": "Ortak",
                                    "label_en": "Partner",
                                },
                            ],
                        ),
                        BriefField(
                            key="partners",
                            label_tr="Ortaklar (her satıra bir tane: İsim, Ülke, Tip)",
                            label_en="Partners (one per line: Name, Country, Type)",
                            type="textarea",
                            max_length=2000,
                            help_text_tr="Min 3 ortak, 3 farklı AB MS/AC ülkesinden zorunlu",
                            help_text_en=("Min 3 partners from 3 different EU MS/AC required"),
                        ),
                    ],
                ),
                BriefSection(
                    title_tr="Etki ve Bütçe",
                    title_en="Impact & Budget",
                    fields=[
                        BriefField(
                            key="target_users",
                            label_tr="Hedef kullanıcılar/paydaşlar",
                            label_en="Target users/stakeholders",
                            type="textarea",
                            max_length=1500,
                        ),
                        BriefField(
                            key="market_size",
                            label_tr="Pazar büyüklüğü (TAM/SAM/SOM)",
                            label_en="Market size (TAM/SAM/SOM)",
                            type="textarea",
                            max_length=1000,
                            required=False,
                        ),
                        BriefField(
                            key="budget_request_eur",
                            label_tr="Talep edilen bütçe (EUR)",
                            label_en="Requested budget (EUR)",
                            type="currency",
                        ),
                        BriefField(
                            key="duration_months",
                            label_tr="Proje süresi (ay)",
                            label_en="Project duration (months)",
                            type="number",
                        ),
                    ],
                ),
            ]
        )

    def get_template_path(self) -> str:
        return str(_TEMPLATE_PATH)

    def get_prompt_path(self, agent_id: str) -> str:
        return str(_PROMPTS_DIR / agent_id / "v1.md")

    def validate_draft(
        self, draft: dict[str, Any], metadata: CallMetadata
    ) -> list[ValidationIssue]:
        issues: list[ValidationIssue] = []
        del metadata  # not needed for this ruleset

        # Page limit (heuristic: ~500 words/page).
        total_words = sum(
            len(str(draft.get(key) or "").split())
            for key in ("excellence_md", "impact_md", "implementation_md")
        )
        estimated_pages = max(1, total_words // self.words_per_page)
        if estimated_pages > self.page_limit:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section=None,
                    code="page_limit_exceeded",
                    message_tr=(
                        f"Tahmini toplam sayfa sayısı {estimated_pages}, "
                        f"limit {self.page_limit}."
                    ),
                    message_en=(
                        f"Estimated total pages {estimated_pages}, " f"limit is {self.page_limit}."
                    ),
                    suggestion=("Trim Excellence section first — typically over-written."),
                )
            )

        # Required subsections — heading prefix must appear in the
        # corresponding markdown body.
        for section, prefixes in _SUBSECTION_PREFIXES.items():
            section_md = str(draft.get(f"{section}_md") or "")
            for prefix in prefixes:
                if f"## {prefix}" not in section_md:
                    issues.append(
                        ValidationIssue(
                            severity="blocker",
                            section=section,
                            code="missing_subsection",
                            message_tr=f"{prefix} alt-bölümü eksik.",
                            message_en=f"{prefix} subsection missing.",
                        )
                    )

        # Consortium sufficiency — pulled from brief.partners (the user
        # field), not from the generated text. Each non-empty line is
        # one partner.
        partners_blob = str(draft.get("brief", {}).get("partners") or "")
        partner_lines = [line for line in partners_blob.splitlines() if line.strip()]
        if 0 < len(partner_lines) < self.min_partners:
            # Only flag if the user provided partners at all — empty
            # blob means brief is incomplete, separate concern.
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="implementation",
                    code="insufficient_consortium",
                    message_tr=(
                        f"Min {self.min_partners} ortak gerekli "
                        f"({self.min_countries} farklı AB MS/AC ülkesinden)."
                    ),
                    message_en=(
                        f"Minimum {self.min_partners} partners required "
                        f"(from {self.min_countries} different EU MS/AC)."
                    ),
                )
            )

        # DNSH — recommended for Cluster 4, formally checked by
        # ComplianceReviewer (S2) for clusters that require it.
        impact_md = str(draft.get("impact_md") or "")
        if "DNSH" not in impact_md and "do no significant harm" not in impact_md.lower():
            issues.append(
                ValidationIssue(
                    severity="warning",
                    section="impact",
                    code="missing_dnsh",
                    message_tr=(
                        "DNSH (Do No Significant Harm) bahsi yok. "
                        "Cluster 4'te zorunlu değil ama önerilir."
                    ),
                    message_en=(
                        "DNSH not mentioned. Not mandatory for Cluster 4 " "but recommended."
                    ),
                )
            )

        # Gender dimension — HE evaluation criterion.
        combined = (
            str(draft.get("excellence_md") or "") + str(draft.get("impact_md") or "")
        ).lower()
        if "gender" not in combined:
            issues.append(
                ValidationIssue(
                    severity="warning",
                    section="excellence",
                    code="missing_gender_dimension",
                    message_tr=(
                        "Gender dimension bahsi yok. HE evaluation " "criteria gereği önerilir."
                    ),
                    message_en=(
                        "Gender dimension not addressed. Recommended " "per HE evaluation criteria."
                    ),
                )
            )

        return issues

    def export_docx(self, proposal: dict[str, Any]) -> bytes:
        if not _TEMPLATE_PATH.is_file():
            raise FileNotFoundError(
                f"HE RIA template missing: {_TEMPLATE_PATH}. "
                "Run `python -m src.programs.horizon_eu_ria.template_builder` to regenerate."
            )

        doc = Document(str(_TEMPLATE_PATH))
        _strip_template_body(doc)
        _set_proposal_title(
            doc,
            title=str(proposal.get("title") or "Untitled Project"),
            acronym=str(proposal.get("acronym") or ""),
        )

        draft = proposal.get("draft") or {}
        for source_key, header in _SECTION_HEADERS:
            doc.add_heading(header, level=1)
            body = str(draft.get(source_key) or "").strip()
            if body:
                render_markdown(doc, body)
            else:
                doc.add_paragraph(f"[{source_key} not yet generated.]")

        # AI disclosure section — Compliance Reviewer fills the body
        # post-generation; for now we emit the placeholder heading so
        # the ToC (when generated) shows the slot.
        doc.add_heading("Use of AI Tools in This Proposal", level=1)
        ai_text = str(proposal.get("ai_disclosure_text") or "").strip()
        if ai_text:
            doc.add_paragraph(ai_text)
        else:
            doc.add_paragraph(
                "[Auto-generated by ComplianceReviewer based on per-sentence "
                "provenance metadata.]"
            ).runs[0].italic = True

        # References — Hallucination Hunter populates after verification.
        doc.add_heading("References", level=1)
        bibliography = proposal.get("bibliography") or []
        if bibliography:
            for idx, entry in enumerate(bibliography, start=1):
                doc.add_paragraph(f"[{idx}] {entry}", style="List Number")
        else:
            doc.add_paragraph("[References numbered after citation verification.]").runs[
                0
            ].italic = True

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_xlsx_budget(self, proposal: dict[str, Any]) -> bytes | None:
        # Lump Sum Excel export lands in S2.D8 (next task per the spec).
        del proposal
        return None


# ── Helpers ────────────────────────────────────────────────────────────


_TEMPLATE_LEADING_PARAGRAPHS_KEPT = 4  # title, subtitle, project title, acronym


def _strip_template_body(doc: DocumentT) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    paragraphs_seen = 0
    for child in children:
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            paragraphs_seen += 1
            if paragraphs_seen <= _TEMPLATE_LEADING_PARAGRAPHS_KEPT:
                continue
        if tag == "sectPr":
            continue
        body.remove(child)


def _set_proposal_title(doc: DocumentT, *, title: str, acronym: str) -> None:
    for para in doc.paragraphs:
        if "[project_title]" in para.text:
            for run in para.runs:
                run.text = run.text.replace("[project_title]", title)
        if "[acronym]" in para.text:
            for run in para.runs:
                run.text = run.text.replace("[acronym]", acronym or "—")


__all__ = ["HorizonEURIAModule"]


# Suppress unused import warning when Pt is referenced only in helpers.
_ = Pt
