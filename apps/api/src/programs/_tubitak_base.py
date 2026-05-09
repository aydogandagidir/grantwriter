"""Shared base class for TÜBİTAK industrial-R&D programmes.

Per docs/07 §5: TÜBİTAK 1501 (Sanayi AR-GE) and 1507 (KOBİ AR-GE
Başlangıç) share ~80% of their structure — same AGY100 form, same
section layout (B1-B4 / C1-C3 / D1-D4), same fatura-bazlı budget
shape. They differ in duration limits, budget caps, and the KOBİ
requirement. This module is the common floor.

Subclasses provide:
- ``program_id``, ``name_tr``, ``name_en`` (always)
- ``get_template_path`` (template file location)
- ``get_prompt_path`` (per-agent prompt file)
- Optional overrides: ``duration_min_months`` / ``duration_max_months``
  / ``b2_min_words`` / ``requires_kobi`` / ``budget_max_tl``

Validation rules are driven by the class-level flags, so subclasses
usually don't override ``validate_draft`` — they just set parameters
and inherit the rule logic.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any, ClassVar

from docx import Document
from docx.document import Document as DocumentT
from docx.shared import Pt
from pydantic import BaseModel, ConfigDict

from src.exports.md_to_docx import render_markdown
from src.programs.base import (
    BaseProgramModule,
    BriefField,
    BriefSchema,
    BriefSection,
    CallMetadata,
    ProgrammeLanguage,
    ValidationIssue,
)

_SUBSECTION_MAP: dict[str, list[str]] = {
    "excellence": [
        "B1_proje_konusu_ve_amaclari",
        "B2_yenilikci_yonleri",
        "B3_yontem_ve_teknik",
        "B4_literature_review",
    ],
    "impact": [
        "C1_ekonomik_ve_ulusal_kazanim",
        "C2_yaygin_etki",
        "C3_pazar_analizi",
    ],
    "implementation": [
        "D1_is_paketleri",
        "D2_zaman_planlamasi",
        "D3_butce",
        "D4_proje_yonetimi_ve_riskler",
    ],
}


_SECTION_HEADERS: list[tuple[str, str, str]] = [
    # (markdown source key on proposal["draft"], DOCX heading text, log label)
    ("excellence_md", "B. Bilimsel ve Teknolojik Detaylar", "excellence"),
    ("impact_md", "C. Yaygın Etki ve Pazar", "impact"),
    ("implementation_md", "D. Uygulama", "implementation"),
]

_BUDGET_CATEGORIES = ["M1", "M2", "M3", "M4", "M5", "M6"]


# Subsection labels for the PRODİS field-by-field copy view.
# Order matters — this is the order the user copies into TÜBİTAK's
# online portal.
class ProdisField(BaseModel):
    """One row in the PRODİS field-by-field copy view.

    The TÜBİTAK PRODİS portal has no public API; users paste the
    proposal section by section. This view gives them the per-field
    plain text + a copy button (frontend) per docs/07 §4.3.
    """

    model_config = ConfigDict(frozen=True)

    key: str
    """Subsection identifier — matches the keys in
    :data:`TUBITAKBaseModule.subsection_map`."""
    label_tr: str
    label_en: str
    value: str
    """Plain text (markdown stripped) for direct paste into PRODİS."""


PRODIS_FIELD_LABELS: list[tuple[str, str, str, str, str]] = [
    # (subsection_id, prefix, source_md_key, label_tr, label_en)
    ("B1_proje_konusu_ve_amaclari", "B1", "excellence_md",
     "B1 Proje Konusu ve Amaçları", "B1 Project Subject and Objectives"),
    ("B2_yenilikci_yonleri", "B2", "excellence_md",
     "B2 Yenilikçi Yönler", "B2 Innovative Aspects"),
    ("B3_yontem_ve_teknik", "B3", "excellence_md",
     "B3 Yöntem ve Teknik", "B3 Method and Technique"),
    ("B4_literature_review", "B4", "excellence_md",
     "B4 Literatür Taraması", "B4 Literature Review"),
    ("C1_ekonomik_ve_ulusal_kazanim", "C1", "impact_md",
     "C1 Ekonomik ve Ulusal Kazanım", "C1 Economic and National Gain"),
    ("C2_yaygin_etki", "C2", "impact_md",
     "C2 Yaygın Etki", "C2 Widespread Impact"),
    ("C3_pazar_analizi", "C3", "impact_md",
     "C3 Pazar Analizi", "C3 Market Analysis"),
    ("D1_is_paketleri", "D1", "implementation_md",
     "D1 İş Paketleri", "D1 Work Packages"),
    ("D2_zaman_planlamasi", "D2", "implementation_md",
     "D2 Zaman Planlaması", "D2 Time Planning"),
    ("D3_butce", "D3", "implementation_md",
     "D3 Bütçe", "D3 Budget"),
    ("D4_proje_yonetimi_ve_riskler", "D4", "implementation_md",
     "D4 Proje Yönetimi ve Riskler", "D4 Project Management and Risks"),
]

_TEMPLATE_LEADING_PARAGRAPHS_KEPT = 3
"""How many leading template paragraphs to preserve when stripping the
placeholder body (cover heading + subtitle + title line)."""


class TUBITAKBaseModule(BaseProgramModule):
    """Common ground for TÜBİTAK programmes (1501, 1507, ...).

    Subclasses set class-level parameters (duration, KOBİ, budget cap)
    and provide the per-programme template + prompts paths. The brief
    schema, call parsing, validation rule engine, and DOCX export are
    inherited unchanged in the common case.
    """

    funder: ClassVar[str] = "TÜBİTAK"
    language: ClassVar[ProgrammeLanguage] = "tr"

    sections: ClassVar[list[str]] = ["excellence", "impact", "implementation"]
    subsection_map: ClassVar[dict[str, list[str]]] = _SUBSECTION_MAP

    # Defaults aligned with 1501 — 1507 / future TÜBİTAK programmes
    # override as needed.
    duration_min_months: ClassVar[int] = 18
    duration_max_months: ClassVar[int] = 36
    requires_kobi: ClassVar[bool] = False
    budget_max_tl: ClassVar[int | None] = None
    b2_min_words: ClassVar[int] = 800

    # ── Plugin API ─────────────────────────────────────────────────────

    def parse_call(self, call_text: str, call_metadata: dict[str, Any]) -> CallMetadata:
        del call_text  # raw text not needed — Call Analyst already parsed
        return CallMetadata.model_validate({**call_metadata, "language": "tr"})

    def get_brief_schema(self) -> BriefSchema:
        """Default TÜBİTAK brief — title, problem, solution, duration.

        Full AGY100 has 28 fields; the writer agents only consume a
        handful for v1. Subclasses can override to add KOBİ flag,
        budget breakdowns, etc. when their writers need more.
        """

        return BriefSchema(
            sections=[
                BriefSection(
                    title_tr="Proje Özü",
                    title_en="Project Core",
                    fields=[
                        BriefField(
                            key="title",
                            label_tr="Proje başlığı",
                            label_en="Project title",
                            type="text",
                            max_length=200,
                        ),
                        BriefField(
                            key="problem_statement",
                            label_tr="Çözmeye çalıştığınız problem",
                            label_en="Problem you're solving",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="proposed_solution",
                            label_tr="Önerdiğiniz çözüm",
                            label_en="Your proposed solution",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="duration_months",
                            label_tr="Proje süresi (ay)",
                            label_en="Project duration (months)",
                            type="number",
                        ),
                    ],
                ),
            ]
        )

    def validate_draft(
        self, draft: dict[str, Any], metadata: CallMetadata
    ) -> list[ValidationIssue]:
        """Apply the TÜBİTAK rule set, parametrised by class-level flags.

        Always-on:
        - B2 minimum word count (TÜBİTAK panel rejection trigger #1)
        - Project duration range

        Conditional:
        - If ``requires_kobi`` is True → blocker when brief.is_sme is False
        - If ``budget_max_tl`` is set → blocker when budget exceeds it
        """

        del metadata  # unused in v1; kept for plugin parity
        issues: list[ValidationIssue] = []

        excellence_md = str(draft.get("excellence_md") or "")
        b2_text = _extract_subsection_body(excellence_md, "B2")
        if len(b2_text.split()) < self.b2_min_words:
            issues.append(self._issue_b2_too_short())

        duration = int(draft.get("brief", {}).get("duration_months") or 0)
        if duration and not (self.duration_min_months <= duration <= self.duration_max_months):
            issues.append(self._issue_invalid_duration(duration))

        if self.requires_kobi:
            issue = self._check_kobi(draft)
            if issue is not None:
                issues.append(issue)

        if self.budget_max_tl is not None:
            issue = self._check_budget_cap(draft)
            if issue is not None:
                issues.append(issue)

        return issues

    def export_docx(self, proposal: dict[str, Any]) -> bytes:
        """Render the AGY100-shaped DOCX from the proposal dict.

        Uses the subclass's ``get_template_path`` for the actual file —
        each programme can ship its own template even if the layout is
        the same.
        """

        from pathlib import Path

        template_path = Path(self.get_template_path())
        if not template_path.is_file():
            raise FileNotFoundError(
                f"Template missing: {template_path}. Run the programme's "
                f"template_builder to regenerate."
            )

        doc = Document(str(template_path))
        _strip_template_body(doc)
        _set_proposal_title(doc, str(proposal.get("title") or "İsimsiz Proje"))

        draft = proposal.get("draft") or {}
        for source_key, header, _ in _SECTION_HEADERS:
            doc.add_heading(header, level=1)
            body = str(draft.get(source_key) or "").strip()
            if body:
                render_markdown(doc, body)
            else:
                doc.add_paragraph(f"[{source_key} bölümü henüz üretilmedi.]")

        budget = (proposal.get("budget") or {}).get("by_category") or {}
        _render_budget_table(doc, budget)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_xlsx_budget(self, proposal: dict[str, Any]) -> bytes | None:
        """Excel budget export lands in S2.D8 (per docs/07 §4.4)."""

        del proposal
        return None

    def get_prodis_fields(self, proposal: dict[str, Any]) -> list[ProdisField]:
        """Render the proposal as 11 PRODİS-portal-ready text fields.

        TÜBİTAK has no public API; users copy each section manually.
        This method splits the writers' Markdown into per-subsection
        plain text (no headings, no formatting characters), so the
        frontend can show one "Copy" button per field.

        Order matches :data:`PRODIS_FIELD_LABELS` — the same order the
        PRODİS portal renders.
        """

        draft = proposal.get("draft") or {}
        fields: list[ProdisField] = []
        for subsection_id, prefix, md_key, label_tr, label_en in PRODIS_FIELD_LABELS:
            section_md = str(draft.get(md_key) or "")
            body_md = _extract_subsection_body(section_md, prefix)
            value = _markdown_to_plain_text(body_md)
            fields.append(
                ProdisField(
                    key=subsection_id,
                    label_tr=label_tr,
                    label_en=label_en,
                    value=value,
                )
            )
        return fields

    # ── Issue builders (overridable per programme for wording) ────────

    def _issue_b2_too_short(self) -> ValidationIssue:
        return ValidationIssue(
            severity="blocker",
            section="excellence",
            code="b2_too_short",
            message_tr=(
                f"B2 (Yenilikçi Yönler) en az {self.b2_min_words} kelime olmalı. "
                "TÜBİTAK panel reddetme tetikleyicisi #1."
            ),
            message_en=(
                f"B2 (Innovative Aspects) must be at least {self.b2_min_words} words. "
                "Most critical section in TÜBİTAK eval."
            ),
            suggestion=(
                "Mevcut state-of-the-art'ı detaylandırın, sizin "
                "yaklaşımınızın farkını net belirtin."
            ),
        )

    def _issue_invalid_duration(self, duration: int) -> ValidationIssue:
        return ValidationIssue(
            severity="blocker",
            section="implementation",
            code="invalid_duration",
            message_tr=(
                f"Proje süresi {duration} ay. {self.name_tr}'de "
                f"{self.duration_min_months}-{self.duration_max_months} "
                "ay arası zorunlu."
            ),
            message_en=(
                f"Project duration {duration} months. {self.name_en} requires "
                f"{self.duration_min_months}-{self.duration_max_months} months."
            ),
        )

    def _check_kobi(self, draft: dict[str, Any]) -> ValidationIssue | None:
        """Return a ``not_kobi`` blocker if the brief flags non-SME company.

        The brief's ``is_sme`` field is a boolean the user fills in (or
        derived from headcount/revenue). Default-True so missing data
        doesn't false-positive — KOBİ is a self-attestation in TÜBİTAK
        applications anyway.
        """

        brief = draft.get("brief") or {}
        is_sme = brief.get("is_sme")
        # Treat "missing" as SME=true (pessimistic flagging would
        # block every freshly-created proposal).
        if is_sme is False:
            return ValidationIssue(
                severity="blocker",
                section="implementation",
                code="not_kobi",
                message_tr=(
                    f"{self.name_tr} sadece KOBİ'lere açık. "
                    "Şirket KOBİ tanımına uymuyor (250+ personel veya 40M TL+ ciro)."
                ),
                message_en=(
                    f"{self.name_en} accepts only SMEs. "
                    "The company exceeds the SME threshold (>250 employees or >40M TL revenue)."
                ),
            )
        return None

    def _check_budget_cap(self, draft: dict[str, Any]) -> ValidationIssue | None:
        """Return a ``budget_exceeded`` blocker if total > programme cap."""

        budget = draft.get("budget") or {}
        total = budget.get("total_tl")
        try:
            total_value = float(total) if total is not None else 0.0
        except (TypeError, ValueError):
            return None

        cap = self.budget_max_tl
        if cap is None or total_value <= cap:
            return None

        return ValidationIssue(
            severity="blocker",
            section="implementation",
            code="budget_exceeded",
            message_tr=(
                f"Bütçe {total_value:,.0f} TL. "
                f"{self.name_tr} limiti {cap:,.0f} TL."
            ),
            message_en=(
                f"Budget {total_value:,.0f} TL exceeds "
                f"{self.name_en} cap of {cap:,.0f} TL."
            ),
        )


# ── Helpers (module-level so subclasses and tests can import) ─────────


def _strip_template_body(doc: DocumentT) -> None:
    """Remove placeholder paragraphs / tables the template builder seeded.

    Cover heading + subtitle + proposal-title-line are kept (the first
    three paragraphs); everything after is dropped so the export
    overwrites rather than appending below.
    """

    body = doc.element.body
    children = list(body.iterchildren())
    paragraphs_seen = 0
    for child in children:
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            paragraphs_seen += 1
            if paragraphs_seen <= _TEMPLATE_LEADING_PARAGRAPHS_KEPT:
                continue
        if tag == "sectPr":
            # Page-section properties live at the end; preserve.
            continue
        body.remove(child)


def _set_proposal_title(doc: DocumentT, title: str) -> None:
    """Replace the ``[proje_basligi]`` placeholder with the real title."""

    for para in doc.paragraphs:
        if "[proje_basligi]" in para.text:
            for run in para.runs:
                if "[proje_basligi]" in run.text:
                    run.text = run.text.replace("[proje_basligi]", title)
            return


def _render_budget_table(
    doc: DocumentT, by_category: dict[str, list[dict[str, Any]]]
) -> None:
    doc.add_heading("Bütçe Tablosu", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    for idx, label in enumerate(
        ["Gider Kalemi", "Açıklama", "Miktar", "Birim Fiyat (TL)", "Toplam (TL)", "Ay"]
    ):
        headers[idx].text = label
        run = headers[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    grand_total = 0.0
    for category in _BUDGET_CATEGORIES:
        items = by_category.get(category) or []
        for item in items:
            row = table.add_row().cells
            row[0].text = category
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("quantity", ""))
            row[3].text = _fmt_money(item.get("unit_price"))
            total = float(item.get("total") or 0)
            row[4].text = _fmt_money(total)
            row[5].text = str(item.get("month", ""))
            grand_total += total

    if grand_total > 0:
        total_row = table.add_row().cells
        total_row[0].text = "TOPLAM"
        total_row[0].paragraphs[0].runs[0].bold = True
        total_row[4].text = _fmt_money(grand_total)
        total_row[4].paragraphs[0].runs[0].bold = True


def _fmt_money(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return str(value)


def _extract_subsection_body(md: str, prefix: str) -> str:
    """Body of a ``## <prefix> ...`` heading up to the next ``## `` or EOF."""

    pattern = re.compile(
        rf"^##\s+{re.escape(prefix)}\b.*$",
        re.MULTILINE,
    )
    match = pattern.search(md)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^##\s+", md[start:], re.MULTILINE)
    end = start + next_match.start() if next_match else len(md)
    return md[start:end].strip()


_MD_HEADING_RE = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_MD_BOLD_ITALIC_RE = re.compile(r"\*{1,3}([^*]+?)\*{1,3}")
_MD_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_MD_LIST_BULLET_RE = re.compile(r"^[\s]*[-*+]\s+", re.MULTILINE)
_MD_LIST_ORDERED_RE = re.compile(r"^[\s]*\d+\.\s+", re.MULTILINE)
_MD_BLOCKQUOTE_RE = re.compile(r"^>\s+", re.MULTILINE)
_MULTI_BLANK_RE = re.compile(r"\n{3,}")


def _markdown_to_plain_text(md: str) -> str:
    """Strip Markdown syntax for the PRODİS copy-paste view.

    Not a full Markdown parser — handles the subset our writer agents
    actually emit (headings, bold, italic, inline code, links, ordered
    and unordered lists, blockquotes). The output is plain text the
    user pastes into TÜBİTAK PRODİS, which doesn't render Markdown.

    Tables intentionally pass through with their pipe characters
    intact — PRODİS doesn't render them, but the user can hand-format
    them in the portal more easily than re-deriving structure from
    stripped text.
    """

    if not md:
        return ""

    text = md
    # Remove heading hashes (keep the heading text — useful in nested
    # subsections within a copied section).
    text = _MD_HEADING_RE.sub("", text)
    # Bold / italic / strikethrough — keep inner text.
    text = _MD_BOLD_ITALIC_RE.sub(r"\1", text)
    # Inline code — keep content, drop backticks.
    text = _MD_INLINE_CODE_RE.sub(r"\1", text)
    # Links — keep label, drop URL.
    text = _MD_LINK_RE.sub(r"\1", text)
    # List bullets / numbers — replace with a clean dash so the
    # paste preserves visual list structure.
    text = _MD_LIST_BULLET_RE.sub("- ", text)
    text = _MD_LIST_ORDERED_RE.sub("- ", text)
    # Blockquotes.
    text = _MD_BLOCKQUOTE_RE.sub("", text)
    # Collapse runs of 3+ blank lines to 2 (paste doesn't need huge gaps).
    text = _MULTI_BLANK_RE.sub("\n\n", text)
    return text.strip()


# ── Shared template builder ───────────────────────────────────────────


# Section structure mirrors docs/07 §4.2 — TÜBİTAK AGY100. Identical
# across 1501 and 1507; if a future programme diverges, override
# ``TEMPLATE_SECTIONS`` per template_builder.
TEMPLATE_SECTIONS: list[tuple[str, list[tuple[str, str]]]] = [
    (
        "B. Bilimsel ve Teknolojik Detaylar",
        [
            ("B1", "Proje Konusu ve Amaçları"),
            ("B2", "Yenilikçi Yönler"),
            ("B3", "Yöntem ve Teknik"),
            ("B4", "Literatür Taraması"),
        ],
    ),
    (
        "C. Yaygın Etki ve Pazar",
        [
            ("C1", "Ekonomik ve Ulusal Kazanım"),
            ("C2", "Yaygın Etki"),
            ("C3", "Pazar Analizi"),
        ],
    ),
    (
        "D. Uygulama",
        [
            ("D1", "İş Paketleri"),
            ("D2", "Zaman Planlaması"),
            ("D3", "Bütçe"),
            ("D4", "Proje Yönetimi ve Riskler"),
        ],
    ),
]

TEMPLATE_PLACEHOLDER_TEXT = "[Bu bölüm proje üretimi sırasında doldurulacaktır.]"


def build_tubitak_template(
    *,
    programme_name: str,
    programme_form: str = "AGY100 — Proje Öneri Bilgileri Formu",
) -> DocumentT:
    """Build a fresh AGY100 scaffold for a TÜBİTAK programme.

    Common shape across 1501, 1507, and future industrial-R&D
    programmes: cover heading, AGY100 subtitle, proposal-title
    placeholder, B/C/D section scaffolding, budget table.

    The runtime export path opens the committed binary template, so
    rebuilding via this helper only matters when the section structure
    or boilerplate copy changes.
    """

    from docx.shared import Cm

    doc = Document()

    # Page margins — TÜBİTAK suggests 2.5 cm all round.
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    title = doc.add_heading(programme_name, level=0)
    for run in title.runs:
        run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run(programme_form).italic = True

    title_p = doc.add_paragraph()
    title_p.add_run("Proje Başlığı: ").bold = True
    title_p.add_run("[proje_basligi]")

    for section_title, subsections in TEMPLATE_SECTIONS:
        doc.add_heading(section_title, level=1)
        for code, label in subsections:
            doc.add_heading(f"{code} {label}", level=2)
            doc.add_paragraph(TEMPLATE_PLACEHOLDER_TEXT)

    doc.add_heading("Bütçe Tablosu", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for idx, label in enumerate(
        ["Gider Kalemi", "Açıklama", "Miktar", "Birim Fiyat (TL)", "Toplam (TL)", "Ay"]
    ):
        header[idx].text = label
        header[idx].paragraphs[0].runs[0].bold = True
    placeholder_row = table.add_row().cells
    placeholder_row[0].text = "[M1-M6]"
    placeholder_row[1].text = "[açıklama]"
    placeholder_row[2].text = "[adet]"
    placeholder_row[3].text = "[birim fiyat]"
    placeholder_row[4].text = "[toplam]"
    placeholder_row[5].text = "[ay]"

    return doc


__all__ = [
    "PRODIS_FIELD_LABELS",
    "ProdisField",
    "TEMPLATE_PLACEHOLDER_TEXT",
    "TEMPLATE_SECTIONS",
    "TUBITAKBaseModule",
    "_BUDGET_CATEGORIES",
    "_SECTION_HEADERS",
    "_SUBSECTION_MAP",
    "_TEMPLATE_LEADING_PARAGRAPHS_KEPT",
    "_extract_subsection_body",
    "_fmt_money",
    "_markdown_to_plain_text",
    "_render_budget_table",
    "_set_proposal_title",
    "_strip_template_body",
    "build_tubitak_template",
]
