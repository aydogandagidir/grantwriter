"""TÜBİTAK 1501 (Sanayi AR-GE Projeleri) programme module.

Implements :class:`~src.programs.base.BaseProgramModule`. S1.D5.T1
ships only the `export_docx` path end-to-end; brief schema, call
parsing, validation, and budget XLSX are minimal placeholders that
later sprints (S2.D7+) flesh out — the contract surface is in place
so the registry, orchestrator, and tests can rely on it.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from docx.document import Document as DocumentT
from docx.shared import Pt

from src.exports.md_to_docx import render_markdown
from src.programs.base import (
    BaseProgramModule,
    BriefField,
    BriefSchema,
    BriefSection,
    CallMetadata,
    ProgrammeLanguage,
    ValidationIssue,
)

_TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "agy100_2026.docx"
_PROMPTS_DIR = Path(__file__).resolve().parent.parent.parent / "agents" / "prompts" / "tubitak_1501"


# Section / subsection layout — must agree with the template_builder output
# and with the Excellence Writer's _SUBSECTION_LAYOUT["tubitak_1501"].
_SUBSECTION_MAP: dict[str, list[str]] = {
    "excellence": [
        "B1_proje_konusu_ve_amaclari",
        "B2_yenilikci_yonleri",
        "B3_yontem_ve_teknik",
        "B4_literature_review",
    ],
    "impact": [
        "C1_ekonomik_ve_ulusal_kazanim",
        "C2_yaygin_etki",
        "C3_pazar_analizi",
    ],
    "implementation": [
        "D1_is_paketleri",
        "D2_zaman_planlamasi",
        "D3_butce",
        "D4_proje_yonetimi_ve_riskler",
    ],
}


_SECTION_HEADERS: list[tuple[str, str, str]] = [
    # (markdown source key on proposal["draft"], DOCX heading text, log label)
    ("excellence_md", "B. Bilimsel ve Teknolojik Detaylar", "excellence"),
    ("impact_md", "C. Yaygın Etki ve Pazar", "impact"),
    ("implementation_md", "D. Uygulama", "implementation"),
]

_BUDGET_CATEGORIES = ["M1", "M2", "M3", "M4", "M5", "M6"]

# How many leading paragraphs of the template scaffold to preserve when
# stripping placeholder body content (cover heading + subtitle + title line).
_TEMPLATE_LEADING_PARAGRAPHS_KEPT = 3


class TUBITAK1501Module(BaseProgramModule):
    program_id: ClassVar[str] = "tubitak_1501"
    name_tr: ClassVar[str] = "TÜBİTAK 1501 Sanayi AR-GE"
    name_en: ClassVar[str] = "TÜBİTAK 1501 Industrial R&D"
    funder: ClassVar[str] = "TÜBİTAK"
    language: ClassVar[ProgrammeLanguage] = "tr"

    sections: ClassVar[list[str]] = ["excellence", "impact", "implementation"]
    subsection_map: ClassVar[dict[str, list[str]]] = _SUBSECTION_MAP

    duration_min_months: int = 18
    duration_max_months: int = 36
    requires_kobi: bool = False  # 1501 accepts large firms; 1507 will override
    b2_min_words: int = 800  # docs/07 §4.5 — TÜBİTAK panel rejection trigger #1

    # ── Plugin API ─────────────────────────────────────────────────────

    def parse_call(self, call_text: str, call_metadata: dict[str, Any]) -> CallMetadata:
        del call_text  # raw text not needed here — Call Analyst already parsed
        return CallMetadata.model_validate({**call_metadata, "language": "tr"})

    def get_brief_schema(self) -> BriefSchema:
        # Minimal shape — full TÜBİTAK form is 28 fields, but for v1 the
        # writer agents only consume a handful. S2.D7 expands.
        return BriefSchema(
            sections=[
                BriefSection(
                    title_tr="Proje Özü",
                    title_en="Project Core",
                    fields=[
                        BriefField(
                            key="title",
                            label_tr="Proje başlığı",
                            label_en="Project title",
                            type="text",
                            max_length=200,
                        ),
                        BriefField(
                            key="problem_statement",
                            label_tr="Çözmeye çalıştığınız problem",
                            label_en="Problem you're solving",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="proposed_solution",
                            label_tr="Önerdiğiniz çözüm",
                            label_en="Your proposed solution",
                            type="textarea",
                            max_length=3000,
                        ),
                        BriefField(
                            key="duration_months",
                            label_tr="Proje süresi (ay)",
                            label_en="Project duration (months)",
                            type="number",
                        ),
                    ],
                ),
            ]
        )

    def get_template_path(self) -> str:
        return str(_TEMPLATE_PATH)

    def get_prompt_path(self, agent_id: str) -> str:
        return str(_PROMPTS_DIR / agent_id / "v1.md")

    def validate_draft(
        self, draft: dict[str, Any], metadata: CallMetadata
    ) -> list[ValidationIssue]:
        # Per docs/07 §4.5: critical TÜBİTAK rules. Full set lands in S2.D7;
        # we ship the rejection-trigger #1 (B2 minimum length) so exports
        # can warn early.
        del metadata  # unused for the v1 ruleset; surfaced for plugin parity
        issues: list[ValidationIssue] = []

        excellence_md = str(draft.get("excellence_md") or "")
        b2_text = _extract_subsection_body(excellence_md, "B2")
        if len(b2_text.split()) < self.b2_min_words:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="excellence",
                    code="b2_too_short",
                    message_tr=(
                        "B2 (Yenilikçi Yönler) en az 800 kelime olmalı. "
                        "TÜBİTAK panel reddetme tetikleyicisi #1."
                    ),
                    message_en=(
                        "B2 (Innovative Aspects) must be at least 800 words. "
                        "Most critical section in TÜBİTAK eval."
                    ),
                    suggestion=(
                        "Mevcut state-of-the-art'ı detaylandırın, sizin "
                        "yaklaşımınızın farkını net belirtin."
                    ),
                )
            )

        duration = int(draft.get("brief", {}).get("duration_months") or 0)
        if duration and not (self.duration_min_months <= duration <= self.duration_max_months):
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section="implementation",
                    code="invalid_duration",
                    message_tr=(
                        f"Proje süresi {duration} ay. TÜBİTAK 1501'de "
                        f"{self.duration_min_months}-{self.duration_max_months} "
                        "ay arası zorunlu."
                    ),
                    message_en=(
                        f"Project duration {duration} months. TÜBİTAK 1501 requires "
                        f"{self.duration_min_months}-{self.duration_max_months} months."
                    ),
                )
            )

        return issues

    def export_docx(self, proposal: dict[str, Any]) -> bytes:
        if not _TEMPLATE_PATH.is_file():
            raise FileNotFoundError(
                f"AGY100 template missing: {_TEMPLATE_PATH}. "
                "Run `python -m src.programs.tubitak_1501.template_builder` to regenerate."
            )

        doc = Document(str(_TEMPLATE_PATH))
        _strip_template_body(doc)
        _set_proposal_title(doc, str(proposal.get("title") or "İsimsiz Proje"))

        draft = proposal.get("draft") or {}
        for source_key, header, _ in _SECTION_HEADERS:
            doc.add_heading(header, level=1)
            body = str(draft.get(source_key) or "").strip()
            if body:
                render_markdown(doc, body)
            else:
                doc.add_paragraph(f"[{source_key} bölümü henüz üretilmedi.]")

        budget = (proposal.get("budget") or {}).get("by_category") or {}
        _render_budget_table(doc, budget)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_xlsx_budget(self, proposal: dict[str, Any]) -> bytes | None:
        # Excel budget export lands in S2.D8 (per docs/07 §4.4). Returning
        # None here is meaningful — the orchestrator skips the XLSX stage.
        del proposal  # unused in v1; kept for interface parity
        return None


# ── Helpers ─────────────────────────────────────────────────────────────


def _strip_template_body(doc: DocumentT) -> None:
    """Remove the placeholder paragraphs / tables the template builder
    seeded so the export overwrites them rather than appending below.

    Cover heading + subtitle + proposal-title-line are kept (they're the
    first three paragraphs); everything after is dropped.
    """

    body = doc.element.body
    children = list(body.iterchildren())
    paragraphs_seen = 0
    for child in children:
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            paragraphs_seen += 1
            if paragraphs_seen <= _TEMPLATE_LEADING_PARAGRAPHS_KEPT:
                continue
        if tag == "sectPr":
            # Page-section properties live at the end; preserve.
            continue
        body.remove(child)


def _set_proposal_title(doc: DocumentT, title: str) -> None:
    """Replace the ``[proje_basligi]`` placeholder with the real title."""

    for para in doc.paragraphs:
        if "[proje_basligi]" in para.text:
            for run in para.runs:
                if "[proje_basligi]" in run.text:
                    run.text = run.text.replace("[proje_basligi]", title)
            return


def _render_budget_table(doc: DocumentT, by_category: dict[str, list[dict[str, Any]]]) -> None:
    doc.add_heading("Bütçe Tablosu", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = table.rows[0].cells
    for idx, label in enumerate(
        ["Gider Kalemi", "Açıklama", "Miktar", "Birim Fiyat (TL)", "Toplam (TL)", "Ay"]
    ):
        headers[idx].text = label
        run = headers[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    grand_total = 0.0
    for category in _BUDGET_CATEGORIES:
        items = by_category.get(category) or []
        for item in items:
            row = table.add_row().cells
            row[0].text = category
            row[1].text = str(item.get("description", ""))
            row[2].text = str(item.get("quantity", ""))
            row[3].text = _fmt_money(item.get("unit_price"))
            total = float(item.get("total") or 0)
            row[4].text = _fmt_money(total)
            row[5].text = str(item.get("month", ""))
            grand_total += total

    if grand_total > 0:
        total_row = table.add_row().cells
        total_row[0].text = "TOPLAM"
        total_row[0].paragraphs[0].runs[0].bold = True
        total_row[4].text = _fmt_money(grand_total)
        total_row[4].paragraphs[0].runs[0].bold = True


def _fmt_money(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return str(value)


def _extract_subsection_body(md: str, prefix: str) -> str:
    """Return the body of a markdown ``## <prefix> ...`` heading up to the
    next ``## `` heading or EOF. Used by validate_draft for B2 length."""

    import re

    pattern = re.compile(
        rf"^##\s+{re.escape(prefix)}\b.*$",
        re.MULTILINE,
    )
    match = pattern.search(md)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^##\s+", md[start:], re.MULTILINE)
    end = start + next_match.start() if next_match else len(md)
    return md[start:end].strip()


__all__ = ["TUBITAK1501Module"]
