"""Build the NLnet (NGI Zero) DOCX scaffold programmatically.

Per docs/07 §7.2. Short form (5-10 pages): three top-level sections
(Project, Team, Budget and Tasks) with brief subheadings. Run
``main()`` once to commit ``templates/nlnet_application.docx``.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentT
from docx.shared import Cm, Pt

PROGRAMME_NAME = "NLnet (NGI Zero) — Cascade Funding Application"
SUBTITLE = "Submitted via NLnet open-call portal"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "Project",
        [
            "Abstract",
            "Main question",
            "Current status",
            "Experience",
            "Comparison to existing solutions",
        ],
    ),
    (
        "Team",
        [
            "Team members",
            "Expertise",
        ],
    ),
    (
        "Budget and Tasks",
        [
            "Tasks and costs",
        ],
    ),
]

PLACEHOLDER_TEXT = "[This section will be populated during draft generation.]"


def build_template() -> DocumentT:
    """Return a fresh NLnet application scaffold."""

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    styles = doc.styles
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    title = doc.add_heading(PROGRAMME_NAME, level=0)
    for run in title.runs:
        run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run(SUBTITLE).italic = True

    title_p = doc.add_paragraph()
    title_p.add_run("Project title: ").bold = True
    title_p.add_run("[project_title]")

    for section_title, subsections in SECTIONS:
        doc.add_heading(section_title, level=1)
        for label in subsections:
            doc.add_heading(label, level=2)
            doc.add_paragraph(PLACEHOLDER_TEXT)

    return doc


def to_bytes() -> bytes:
    buf = BytesIO()
    build_template().save(buf)
    return buf.getvalue()


def main(target: Path | None = None) -> Path:
    target = target or (
        Path(__file__).resolve().parent / "templates" / "nlnet_application.docx"
    )
    target.parent.mkdir(parents=True, exist_ok=True)
    build_template().save(str(target))
    return target


if __name__ == "__main__":
    written = main()
    print(f"wrote: {written}")


__all__ = ["PROGRAMME_NAME", "SECTIONS", "build_template", "main", "to_bytes"]
