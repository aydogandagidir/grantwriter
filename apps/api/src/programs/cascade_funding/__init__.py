"""Cascade Funding (NLnet / NGI Zero) programme module.

Per docs/07 §7. NLnet is the EU's "redistributed" grant mechanism via
NGI (Next Generation Internet) — small grants (€1.5K-€300K), short
applications (5-10 pages). MVP supports the NLnet portal only; other
FSTPs (NGI Sargasso, NGI TALER, etc.) land as sibling sub-modules in
Faz 2 per the directory layout in docs/07 §7.3.

**Design note on sections:** NLnet's form is structured around
``project``, ``team``, ``budget`` rather than HE/TÜBİTAK's
``excellence/impact/implementation``. To stay compatible with the
Saga orchestrator (which dispatches by writer agent ID and stores
results under fixed ``_md`` keys) we KEEP the canonical section names
and map NLnet's subsections beneath them:

  excellence (Project): abstract, main_question, current_status,
                        experience, comparison
  impact (Team):        members, expertise
  implementation (Budget): tasks_and_costs

The DOCX export renders headings as "Project / Team / Budget" so the
final document reads naturally for NLnet evaluators; the storage layer
just sees the standard keys.
"""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any, ClassVar

from docx import Document

from src.exports.md_to_docx import render_markdown
from src.programs.base import (
    BaseProgramModule,
    BriefField,
    BriefSchema,
    BriefSection,
    CallMetadata,
    ProgrammeLanguage,
    ValidationIssue,
)

_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / "templates" / "nlnet_application.docx"
)
_PROMPTS_DIR = (
    Path(__file__).resolve().parent.parent.parent
    / "agents"
    / "prompts"
    / "cascade_funding"
)


_SUBSECTION_MAP: dict[str, list[str]] = {
    "excellence": [
        "abstract",
        "main_question",
        "current_status",
        "experience",
        "comparison",
    ],
    "impact": [
        "team_members",
        "team_expertise",
    ],
    "implementation": [
        "tasks_and_costs",
    ],
}


_SECTION_HEADERS: list[tuple[str, str]] = [
    ("excellence_md", "Project"),
    ("impact_md", "Team"),
    ("implementation_md", "Budget and Tasks"),
]

_TEMPLATE_LEADING_PARAGRAPHS_KEPT = 3


class CascadeFundingModule(BaseProgramModule):
    program_id: ClassVar[str] = "cascade_funding"
    name_tr: ClassVar[str] = "NLnet (NGI Zero) Cascade Funding"
    name_en: ClassVar[str] = "NLnet (NGI Zero) Cascade Funding"
    funder: ClassVar[str] = "NLnet (NGI Zero)"
    language: ClassVar[ProgrammeLanguage] = "en"

    sections: ClassVar[list[str]] = ["excellence", "impact", "implementation"]
    subsection_map: ClassVar[dict[str, list[str]]] = _SUBSECTION_MAP

    page_limit: ClassVar[int] = 10
    words_per_page: ClassVar[int] = 500
    abstract_min_words: ClassVar[int] = 200

    # ── Plugin API ─────────────────────────────────────────────────────

    def parse_call(self, call_text: str, call_metadata: dict[str, Any]) -> CallMetadata:
        del call_text
        merged = {**call_metadata}
        merged.setdefault("language", "en")
        merged.setdefault("page_limit", self.page_limit)
        return CallMetadata.model_validate(merged)

    def get_brief_schema(self) -> BriefSchema:
        return BriefSchema(
            sections=[
                BriefSection(
                    title_tr="Proje Özü",
                    title_en="Project Core",
                    fields=[
                        BriefField(
                            key="title",
                            label_tr="Proje başlığı",
                            label_en="Project title",
                            type="text",
                            max_length=120,
                        ),
                        BriefField(
                            key="main_question",
                            label_tr="Çözmek istediğiniz ana soru",
                            label_en="The main question you want to address",
                            type="textarea",
                            max_length=800,
                        ),
                        BriefField(
                            key="proposed_solution",
                            label_tr="Önerdiğiniz çözüm",
                            label_en="Your proposed solution",
                            type="textarea",
                            max_length=2000,
                        ),
                        BriefField(
                            key="current_status",
                            label_tr="Mevcut durum (varsa)",
                            label_en="Current status of the work (if any)",
                            type="textarea",
                            max_length=1000,
                            required=False,
                        ),
                    ],
                ),
                BriefSection(
                    title_tr="Bütçe",
                    title_en="Budget",
                    fields=[
                        BriefField(
                            key="budget_request_eur",
                            label_tr="Talep edilen bütçe (EUR)",
                            label_en="Requested budget (EUR)",
                            type="currency",
                            help_text_en=(
                                "NLnet typically grants between €1.5K and €50K per "
                                "project; up to €300K for some calls."
                            ),
                        ),
                        BriefField(
                            key="duration_months",
                            label_tr="Proje süresi (ay)",
                            label_en="Project duration (months)",
                            type="number",
                        ),
                    ],
                ),
            ]
        )

    def get_template_path(self) -> str:
        return str(_TEMPLATE_PATH)

    def get_prompt_path(self, agent_id: str) -> str:
        return str(_PROMPTS_DIR / agent_id / "v1.md")

    def validate_draft(
        self, draft: dict[str, Any], metadata: CallMetadata
    ) -> list[ValidationIssue]:
        del metadata
        issues: list[ValidationIssue] = []

        # Page limit (heuristic — NLnet enforces 10 pages strictly).
        total_words = sum(
            len(str(draft.get(key) or "").split())
            for key in ("excellence_md", "impact_md", "implementation_md")
        )
        estimated_pages = max(1, total_words // self.words_per_page)
        if estimated_pages > self.page_limit:
            issues.append(
                ValidationIssue(
                    severity="blocker",
                    section=None,
                    code="page_limit_exceeded",
                    message_tr=(
                        f"Tahmini toplam sayfa {estimated_pages}, "
                        f"NLnet limiti {self.page_limit}."
                    ),
                    message_en=(
                        f"Estimated pages {estimated_pages}, "
                        f"NLnet limit is {self.page_limit}."
                    ),
                    suggestion="Trim the Project section first.",
                )
            )

        # Abstract minimum length — pulled from excellence_md's
        # ``## Abstract`` heading body, if present.
        excellence_md = str(draft.get("excellence_md") or "")
        abstract_text = _extract_subsection_body(excellence_md, "Abstract")
        if abstract_text and len(abstract_text.split()) < self.abstract_min_words:
            issues.append(
                ValidationIssue(
                    severity="warning",
                    section="excellence",
                    code="abstract_too_short",
                    message_tr=(
                        f"Abstract {len(abstract_text.split())} kelime; "
                        f"NLnet en az {self.abstract_min_words} kelime önerir."
                    ),
                    message_en=(
                        f"Abstract is {len(abstract_text.split())} words; "
                        f"NLnet recommends at least {self.abstract_min_words}."
                    ),
                )
            )

        return issues

    def export_docx(self, proposal: dict[str, Any]) -> bytes:
        if not _TEMPLATE_PATH.is_file():
            raise FileNotFoundError(
                f"NLnet template missing: {_TEMPLATE_PATH}. Run "
                "`python -m src.programs.cascade_funding.template_builder` to regenerate."
            )

        doc = Document(str(_TEMPLATE_PATH))
        _strip_template_body(doc)
        _set_proposal_title(doc, str(proposal.get("title") or "Untitled Project"))

        draft = proposal.get("draft") or {}
        for source_key, header in _SECTION_HEADERS:
            doc.add_heading(header, level=1)
            body = str(draft.get(source_key) or "").strip()
            if body:
                render_markdown(doc, body)
            else:
                doc.add_paragraph(f"[{source_key} not yet generated.]")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_xlsx_budget(self, proposal: dict[str, Any]) -> bytes | None:
        # NLnet uses simple task-cost mapping in the form body, no Excel.
        del proposal
        return None


# ── Helpers (NLnet-specific, English-named templates) ─────────────────


def _strip_template_body(doc: Any) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    paragraphs_seen = 0
    for child in children:
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            paragraphs_seen += 1
            if paragraphs_seen <= _TEMPLATE_LEADING_PARAGRAPHS_KEPT:
                continue
        if tag == "sectPr":
            continue
        body.remove(child)


def _set_proposal_title(doc: Any, title: str) -> None:
    for para in doc.paragraphs:
        if "[project_title]" in para.text:
            for run in para.runs:
                if "[project_title]" in run.text:
                    run.text = run.text.replace("[project_title]", title)
            return


def _extract_subsection_body(md: str, prefix: str) -> str:
    """Match a ``## <prefix>`` heading and return the body up to the next ``## ``.

    Used for the abstract length check; ``prefix`` is matched
    case-insensitively so users typing "abstract" or "Abstract" both work.
    """

    pattern = re.compile(
        rf"^##\s+{re.escape(prefix)}\b.*$",
        re.MULTILINE | re.IGNORECASE,
    )
    match = pattern.search(md)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^##\s+", md[start:], re.MULTILINE)
    end = start + next_match.start() if next_match else len(md)
    return md[start:end].strip()


__all__ = ["CascadeFundingModule"]
