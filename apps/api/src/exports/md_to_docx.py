"""Minimal Markdown → python-docx converter.

Handles the subset that grant-proposal sections actually contain:
``# / ## / ###`` headings, blank-line-separated paragraphs, ``**bold**``,
``*italic*``, ``-`` / ``*`` bullet lists, and ``1.`` numbered lists.
Tables, code fences, links, and images are out of scope (per S1.D5.T1
spec) — tables come from structured data, the rest don't appear in
agent output.

Inline formatting is parsed run-by-run inside each paragraph so a
``This is **important**`` becomes three runs, the middle one bold.
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from docx.document import Document

# ── Inline formatting ───────────────────────────────────────────────────

# Match **bold** first (longer), then *italic*. Greedy-aware patterns
# that don't cross newlines.
_BOLD_RE = re.compile(r"\*\*([^*\n]+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")


def _emit_runs(paragraph: Any, text: str) -> None:
    """Add bold/italic-aware runs to a python-docx paragraph.

    Tokenize the line by alternating between bold matches, italic
    matches, and plain text. We replace bold spans with sentinels so
    italic-matching doesn't eat their inner stars.
    """

    # Use a placeholder protocol so a bold span can contain literal '*'
    # without italic re-matching them — extract bold first, italic next.
    bold_spans: list[str] = []

    def _stash_bold(match: re.Match[str]) -> str:
        bold_spans.append(match.group(1))
        return f"\x00BOLD{len(bold_spans) - 1}\x00"

    intermediate = _BOLD_RE.sub(_stash_bold, text)

    italic_spans: list[str] = []

    def _stash_italic(match: re.Match[str]) -> str:
        italic_spans.append(match.group(1))
        return f"\x00ITALIC{len(italic_spans) - 1}\x00"

    intermediate = _ITALIC_RE.sub(_stash_italic, intermediate)

    # Now intermediate contains plain text + sentinels. Walk it and emit
    # runs, decoding sentinels back to their formatted form.
    sentinel_re = re.compile(r"\x00(BOLD|ITALIC)(\d+)\x00")
    pos = 0
    for m in sentinel_re.finditer(intermediate):
        if m.start() > pos:
            paragraph.add_run(intermediate[pos : m.start()])
        kind = m.group(1)
        idx = int(m.group(2))
        run = paragraph.add_run(bold_spans[idx] if kind == "BOLD" else italic_spans[idx])
        if kind == "BOLD":
            run.bold = True
        else:
            run.italic = True
        pos = m.end()
    if pos < len(intermediate):
        paragraph.add_run(intermediate[pos:])


# ── Block parsing ───────────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+?)\s*#*\s*$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")


def render_markdown(doc: Document, markdown: str) -> None:
    """Append rendered Markdown blocks to ``doc`` in order.

    Empty lines separate paragraphs. Consecutive ``- `` lines collapse
    into one bullet list, consecutive ``1. `` / ``2. `` lines into one
    numbered list. Headings always start their own block.
    """

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        h_match = _HEADING_RE.match(stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            heading = doc.add_heading(level=level)
            _emit_runs(heading, text)
            i += 1
            continue

        if _BULLET_RE.match(stripped):
            i = _consume_list(doc, lines, i, _BULLET_RE, style="List Bullet")
            continue

        if _NUMBERED_RE.match(stripped):
            i = _consume_list(doc, lines, i, _NUMBERED_RE, style="List Number")
            continue

        # Paragraph: gather adjacent non-blank lines into one run.
        para_lines: list[str] = []
        while i < len(lines) and lines[i].strip() and not _is_block_starter(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        _emit_runs(para, " ".join(para_lines))


def _is_block_starter(line: str) -> bool:
    stripped = line.strip()
    return bool(
        _HEADING_RE.match(stripped) or _BULLET_RE.match(stripped) or _NUMBERED_RE.match(stripped)
    )


def _consume_list(
    doc: Document,
    lines: list[str],
    start: int,
    pattern: re.Pattern[str],
    *,
    style: str,
) -> int:
    i = start
    while i < len(lines):
        m = pattern.match(lines[i].strip())
        if not m:
            break
        # Numbered pattern has 2 groups; bullet pattern has 1. Take the last.
        text = m.group(m.lastindex or 1)
        para = doc.add_paragraph(style=style)
        _emit_runs(para, text)
        i += 1
    return i


__all__ = ["render_markdown"]
