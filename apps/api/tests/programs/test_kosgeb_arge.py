"""KOSGEB AR-GE / Yenilik module tests.

Plugin contract + KOSGEB-specific blockers (K2 length, KOBİ requirement,
1-year minimum company age, 12-36 month duration). KOSGEB has 4-category
budget table vs TÜBİTAK's 6.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from src.programs import REGISTRY, KOSGEBARGEModule, get_module
from src.programs.base import CallMetadata
from src.programs.kosgeb_arge.template_builder import build_template, to_bytes

# ── Registry + identity ─────────────────────────────────────────────────


def test_module_in_registry() -> None:
    assert "kosgeb_arge" in REGISTRY
    assert isinstance(get_module("kosgeb_arge"), KOSGEBARGEModule)


def test_template_path_exists() -> None:
    module = get_module("kosgeb_arge")
    assert Path(module.get_template_path()).is_file()


def test_prompt_paths_resolve_to_kosgeb_directory() -> None:
    module = get_module("kosgeb_arge")
    for agent_id in ("excellence_writer", "impact_writer", "implementation_writer"):
        path = Path(module.get_prompt_path(agent_id))
        assert "kosgeb_arge" in path.parts
        assert path.is_file(), f"missing prompt: {path}"


# ── Class-level params ──────────────────────────────────────────────────


def test_class_parameters_reflect_kosgeb_rules() -> None:
    module = get_module("kosgeb_arge")
    assert module.duration_min_months == 12
    assert module.duration_max_months == 36
    assert module.requires_kobi is True
    assert module.min_company_age_years == 1
    assert module.k2_min_words == 600


def test_subsection_map_uses_klm_prefixes() -> None:
    module = get_module("kosgeb_arge")
    excellence = module.subsection_map["excellence"]
    assert any(s.startswith("K1") for s in excellence)
    assert any(s.startswith("K2") for s in excellence)
    assert any(s.startswith("K3") for s in excellence)
    assert all(not s.startswith("B") for s in excellence)  # not TÜBİTAK


# ── parse_call ─────────────────────────────────────────────────────────


def test_parse_call_normalises_to_turkish() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("(test)", {})
    assert isinstance(metadata, CallMetadata)
    assert metadata.language == "tr"


# ── Brief schema ───────────────────────────────────────────────────────


def test_brief_schema_includes_company_info_section() -> None:
    module = get_module("kosgeb_arge")
    schema = module.get_brief_schema()
    assert len(schema.sections) >= 2
    field_keys = {f.key for s in schema.sections for f in s.fields}
    assert {"company_age_years", "is_sme"} <= field_keys


# ── validate_draft ─────────────────────────────────────────────────────


def _k2_long_enough() -> str:
    return "## K2 Yenilik Niteliği\n\n" + " ".join(["kelime"] * 650) + "\n"


def test_validate_draft_clean_returns_no_blockers() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _k2_long_enough(),
        "brief": {
            "duration_months": 24,
            "is_sme": True,
            "company_age_years": 3,
        },
    }
    issues = module.validate_draft(draft, metadata)
    assert [i for i in issues if i.severity == "blocker"] == []


def test_validate_draft_blocks_when_k2_too_short() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("", {})
    short_excellence = "## K2 Yenilik Niteliği\n\nÇok kısa.\n"
    issues = module.validate_draft({"excellence_md": short_excellence}, metadata)
    assert any(i.code == "k2_too_short" and i.severity == "blocker" for i in issues)


def test_validate_draft_blocks_when_company_too_young() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _k2_long_enough(),
        "brief": {
            "duration_months": 24,
            "is_sme": True,
            "company_age_years": 0.5,  # 6 months — too young
        },
    }
    issues = module.validate_draft(draft, metadata)
    assert any(
        i.code == "company_too_young" and i.severity == "blocker" for i in issues
    )


def test_validate_draft_blocks_when_not_kobi() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _k2_long_enough(),
        "brief": {"duration_months": 24, "is_sme": False},
    }
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "not_kobi" and i.severity == "blocker" for i in issues)


def test_validate_draft_blocks_when_duration_outside_kosgeb_range() -> None:
    module = get_module("kosgeb_arge")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _k2_long_enough(),
        "brief": {"duration_months": 48, "is_sme": True},
    }
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "invalid_duration" and i.severity == "blocker" for i in issues)


# ── DOCX export ─────────────────────────────────────────────────────────


def test_export_docx_renders_klm_sections() -> None:
    module = get_module("kosgeb_arge")
    proposal = {
        "title": "AKILLI-TEKSTIL",
        "draft": {
            "excellence_md": "## K1 Proje Konusu\n\nKOBİ tekstil için yenilik.\n",
            "impact_md": "## L1 Ekonomik Kazanım\n\nTürkiye ihracat artışı.\n",
            "implementation_md": "## M1 İş-Zaman Planı\n\nWP1, WP2, WP3.\n",
        },
        "budget": {
            "by_category": {
                "Personel": [{"description": "Yazılım mühendisi", "total": 600_000}]
            }
        },
    }
    blob = module.export_docx(proposal)
    doc = Document(BytesIO(blob))

    body = "\n".join(p.text for p in doc.paragraphs)
    assert "AKILLI-TEKSTIL" in body
    assert "K. Proje Tanımı" in body
    assert "L. Etki ve Pazar" in body
    assert "M. Uygulama" in body


def test_export_xlsx_budget_returns_none() -> None:
    module = get_module("kosgeb_arge")
    assert module.export_xlsx_budget({"budget": {}}) is None


# ── Template builder ────────────────────────────────────────────────────


def test_template_builder_smoke() -> None:
    doc = build_template()
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "KOSGEB" in text
    assert "K1" in text and "K2" in text and "K3" in text
    assert "L1" in text and "L2" in text
    assert "M1" in text and "M2" in text and "M3" in text


def test_template_builder_to_bytes_returns_valid_docx() -> None:
    blob = to_bytes()
    doc = Document(BytesIO(blob))
    assert any("KOSGEB" in p.text for p in doc.paragraphs)
