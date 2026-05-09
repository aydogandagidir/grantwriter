"""Cascade Funding (NLnet) module tests.

Plugin contract + NLnet-specific rules: 10-page limit, 200-word abstract
minimum, English language, project/team/budget sections mapped onto the
canonical excellence/impact/implementation keys.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from src.programs import REGISTRY, CascadeFundingModule, get_module
from src.programs.base import CallMetadata
from src.programs.cascade_funding.template_builder import build_template, to_bytes

# ── Registry + identity ─────────────────────────────────────────────────


def test_module_in_registry() -> None:
    assert "cascade_funding" in REGISTRY
    assert isinstance(get_module("cascade_funding"), CascadeFundingModule)


def test_template_path_exists() -> None:
    module = get_module("cascade_funding")
    assert Path(module.get_template_path()).is_file()


def test_prompt_paths_resolve_to_cascade_funding_directory() -> None:
    module = get_module("cascade_funding")
    for agent_id in ("excellence_writer", "impact_writer", "implementation_writer"):
        path = Path(module.get_prompt_path(agent_id))
        assert "cascade_funding" in path.parts
        assert path.is_file(), f"missing prompt: {path}"


# ── Class-level params ──────────────────────────────────────────────────


def test_class_parameters_reflect_nlnet_rules() -> None:
    module = get_module("cascade_funding")
    assert module.page_limit == 10
    assert module.abstract_min_words == 200
    assert module.language == "en"
    assert module.funder.startswith("NLnet")


def test_subsection_map_uses_nlnet_form_names() -> None:
    """Sections stay canonical (excellence/impact/implementation) for
    Saga compat, but subsections name NLnet's actual form fields."""

    module = get_module("cascade_funding")
    assert "abstract" in module.subsection_map["excellence"]
    assert "main_question" in module.subsection_map["excellence"]
    assert "team_members" in module.subsection_map["impact"]
    assert "tasks_and_costs" in module.subsection_map["implementation"]


# ── parse_call ─────────────────────────────────────────────────────────


def test_parse_call_normalises_to_english() -> None:
    module = get_module("cascade_funding")
    metadata = module.parse_call("(test)", {})
    assert isinstance(metadata, CallMetadata)
    assert metadata.language == "en"
    assert metadata.page_limit == 10


# ── Brief schema ───────────────────────────────────────────────────────


def test_brief_schema_has_project_and_budget_sections() -> None:
    module = get_module("cascade_funding")
    schema = module.get_brief_schema()
    titles = {s.title_en for s in schema.sections}
    assert "Project Core" in titles
    assert "Budget" in titles


# ── validate_draft ─────────────────────────────────────────────────────


def _abstract_long_enough() -> str:
    """≥ 200-word abstract body."""

    return "## Abstract\n\n" + " ".join(["word"] * 220) + "\n"


def test_validate_draft_clean_returns_no_blockers() -> None:
    module = get_module("cascade_funding")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _abstract_long_enough(),
        "impact_md": "## Team members\n\nAlice, Bob.\n",
        "implementation_md": "## Tasks and costs\n\n5K, 6 months.\n",
    }
    issues = module.validate_draft(draft, metadata)
    assert [i for i in issues if i.severity == "blocker"] == []


def test_validate_draft_blocks_when_page_limit_exceeded() -> None:
    """10 pages × 500 words = 5000 words. Force 6000 words → 1 page over."""

    module = get_module("cascade_funding")
    metadata = module.parse_call("", {})
    long = " ".join(["word"] * 6000)
    draft = {
        "excellence_md": f"## Abstract\n\n{long}\n",
        "impact_md": "",
        "implementation_md": "",
    }
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "page_limit_exceeded" and i.severity == "blocker" for i in issues)


def test_validate_draft_warns_when_abstract_too_short() -> None:
    module = get_module("cascade_funding")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": "## Abstract\n\nA very short one-sentence abstract.\n",
    }
    issues = module.validate_draft(draft, metadata)
    assert any(
        i.code == "abstract_too_short" and i.severity == "warning" for i in issues
    )


def test_validate_draft_does_not_warn_when_abstract_missing() -> None:
    """No abstract heading means we can't measure it — silent rather
    than false-positive warn."""

    module = get_module("cascade_funding")
    metadata = module.parse_call("", {})
    draft = {"excellence_md": "## Main question\n\nA paragraph.\n"}
    issues = module.validate_draft(draft, metadata)
    assert not any(i.code == "abstract_too_short" for i in issues)


# ── DOCX export ─────────────────────────────────────────────────────────


def test_export_docx_renders_project_team_budget_headings() -> None:
    module = get_module("cascade_funding")
    proposal = {
        "title": "OPEN-AUDIT",
        "draft": {
            "excellence_md": "## Abstract\n\nSecurity audit tooling.\n",
            "impact_md": "## Team members\n\nAlice (maintainer of X).\n",
            "implementation_md": "## Tasks and costs\n\n3 tasks, 6 months.\n",
        },
    }
    blob = module.export_docx(proposal)
    doc = Document(BytesIO(blob))

    headings = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    assert "Project" in headings
    assert "Team" in headings
    assert "Budget and Tasks" in headings

    body = "\n".join(p.text for p in doc.paragraphs)
    assert "OPEN-AUDIT" in body


def test_export_xlsx_budget_returns_none() -> None:
    module = get_module("cascade_funding")
    assert module.export_xlsx_budget({"budget": {}}) is None


# ── Template builder ────────────────────────────────────────────────────


def test_template_builder_smoke() -> None:
    doc = build_template()
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "NLnet" in text
    assert "Project" in text
    assert "Team" in text


def test_template_builder_to_bytes_returns_valid_docx() -> None:
    blob = to_bytes()
    doc = Document(BytesIO(blob))
    headings = [
        p.text
        for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    assert "Abstract" in headings
    assert "Tasks and costs" in headings
