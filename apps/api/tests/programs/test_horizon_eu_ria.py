"""Horizon Europe RIA module tests.

Cover the plugin contract: registry membership, prompt paths, brief
schema shape, validate_draft rules, export_docx output structure.
End-to-end writer-agent run is exercised in
``tests/agents/test_he_writers_e2e.py``.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from src.programs import REGISTRY, HorizonEURIAModule, get_module
from src.programs.base import CallMetadata
from src.programs.horizon_eu_ria.template_builder import (
    AI_DISCLOSURE_BOOKMARK,
    build_template,
    to_bytes,
)

# ── Registry + paths ────────────────────────────────────────────────────


def test_module_in_registry() -> None:
    assert "horizon_eu_ria" in REGISTRY
    assert isinstance(get_module("horizon_eu_ria"), HorizonEURIAModule)


def test_template_path_exists() -> None:
    module = get_module("horizon_eu_ria")
    assert Path(module.get_template_path()).is_file()


def test_prompt_paths_resolve_to_horizon_eu_directory() -> None:
    module = get_module("horizon_eu_ria")
    for agent_id in ("excellence_writer", "impact_writer", "implementation_writer"):
        path = Path(module.get_prompt_path(agent_id))
        # Must resolve under the SHARED horizon_eu directory (not horizon_eu_ria).
        assert "horizon_eu" in path.parts
        assert "horizon_eu_ria" not in path.parts
        assert path.is_file(), f"missing prompt: {path}"


def test_template_builder_produces_valid_docx_with_ai_disclosure_bookmark() -> None:
    blob = to_bytes()
    doc = Document(BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Horizon Europe — Research and Innovation Action" in text
    assert "1.1 Objectives and ambition" in text
    assert "2.1" in text  # Impact subsections
    assert "3.1" in text  # Implementation subsections

    # AI disclosure section heading present + bookmark embedded.
    assert "Use of AI Tools in This Proposal" in text
    raw_xml = doc.element.body.xml
    assert AI_DISCLOSURE_BOOKMARK in raw_xml


# ── Brief schema ────────────────────────────────────────────────────────


def test_brief_schema_has_three_sections_with_required_fields() -> None:
    module = get_module("horizon_eu_ria")
    schema = module.get_brief_schema()
    assert len(schema.sections) == 3
    section_titles_en = [s.title_en for s in schema.sections]
    assert section_titles_en == ["Project Core", "Consortium", "Impact & Budget"]

    field_keys = {f.key for section in schema.sections for f in section.fields}
    # Spot-check the fields docs/07 §3.3 demands.
    expected = {
        "title",
        "acronym",
        "problem_statement",
        "proposed_solution",
        "trl_current",
        "trl_target",
        "role",
        "partners",
        "target_users",
        "budget_request_eur",
        "duration_months",
    }
    assert expected.issubset(field_keys)


# ── parse_call ─────────────────────────────────────────────────────────


def test_parse_call_normalises_language_and_page_limit() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("(call text)", {"scope_summary": "Some scope"})
    assert isinstance(metadata, CallMetadata)
    assert metadata.language == "en"
    assert metadata.page_limit == 45


# ── validate_draft ──────────────────────────────────────────────────────


def _full_draft() -> dict[str, object]:
    """A complete-looking HE draft that should pass every validator."""

    return {
        "excellence_md": (
            "## 1.1 Objectives and ambition\n\nA concrete objective.\n\n"
            "## 1.2 Methodology\n\nGender dimension is integrated.\n\n"
            "## 1.3 State of the art\n\nReferences here.\n\n"
            "## 1.4 Open science practices\n\nFAIR + DMP.\n"
        ),
        "impact_md": (
            "## 2.1 Project's pathways towards impact\n\nKIPs detailed.\n\n"
            "## 2.2 Measures to maximise impact\n\nDNSH considered.\n\n"
            "## 2.3 Summary canvas (key impact pathways)\n\nCanvas table.\n"
        ),
        "implementation_md": (
            "## 3.1 Work plan and resources\n\nWPs.\n\n"
            "## 3.2 Capacity of the participants\n\nTeam.\n\n"
            "## 3.3 Consortium as a whole\n\nGovernance.\n"
        ),
        "brief": {
            "partners": "ACME (NL, SME)\nBeta (DE, RTO)\nGamma (TR, SME)\n",
        },
    }


def test_validate_draft_clean_returns_no_blockers() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})
    issues = module.validate_draft(_full_draft(), metadata)
    blockers = [i for i in issues if i.severity == "blocker"]
    assert blockers == []


def test_validate_draft_flags_missing_subsection() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})

    bad = _full_draft()
    bad["excellence_md"] = (
        "## 1.1 Objectives\n\n.\n\n## 1.3 State of the art\n\n.\n\n" "## 1.4 Open science\n\n.\n"
    )  # 1.2 missing
    issues = module.validate_draft(bad, metadata)
    codes = [i.code for i in issues]
    assert "missing_subsection" in codes
    blocker_messages = [i.message_en for i in issues if i.code == "missing_subsection"]
    assert any("1.2" in msg for msg in blocker_messages)


def test_validate_draft_flags_page_limit_exceeded() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})

    bad = _full_draft()
    long_words = " ".join(["word"] * 30_000)  # ~60 pages at 500 words/page
    bad["excellence_md"] = (
        f"## 1.1 Objectives\n\n{long_words}\n\n"
        "## 1.2 Methodology\n\n.\n\n"
        "## 1.3 State of the art\n\n.\n\n"
        "## 1.4 Open science\n\n.\n"
    )
    issues = module.validate_draft(bad, metadata)
    assert any(i.code == "page_limit_exceeded" and i.severity == "blocker" for i in issues)


def test_validate_draft_warns_when_dnsh_missing() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})

    draft = _full_draft()
    draft["impact_md"] = (
        "## 2.1 Project's pathways towards impact\n\n.\n\n"
        "## 2.2 Measures to maximise impact\n\n.\n\n"  # no DNSH mention
        "## 2.3 Summary canvas (key impact pathways)\n\n.\n"
    )
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "missing_dnsh" and i.severity == "warning" for i in issues)


def test_validate_draft_warns_when_gender_missing() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})

    draft = _full_draft()
    # No "gender" mention anywhere.
    draft["excellence_md"] = (
        "## 1.1 Objectives\n\n.\n\n"
        "## 1.2 Methodology\n\n.\n\n"
        "## 1.3 State of the art\n\n.\n\n"
        "## 1.4 Open science\n\n.\n"
    )
    draft["impact_md"] = (
        "## 2.1 Project's pathways towards impact\n\nDNSH considered.\n\n"
        "## 2.2 Measures to maximise impact\n\n.\n\n"
        "## 2.3 Summary canvas (key impact pathways)\n\n.\n"
    )
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "missing_gender_dimension" and i.severity == "warning" for i in issues)


def test_validate_draft_flags_consortium_too_small() -> None:
    module = get_module("horizon_eu_ria")
    metadata = module.parse_call("", {})

    draft = _full_draft()
    draft["brief"] = {"partners": "ACME (NL, SME)\n"}  # only one partner
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "insufficient_consortium" and i.severity == "blocker" for i in issues)


# ── export_docx ─────────────────────────────────────────────────────────


def test_export_docx_contains_all_sections() -> None:
    module = get_module("horizon_eu_ria")
    proposal = {
        "title": "GREENMOBILITY",
        "acronym": "GMOB",
        "draft": _full_draft(),
        "ai_disclosure_text": (
            "Of 412 sentences in the proposal: 89 were authored by the applicant; "
            "287 were AI-generated; 36 were AI-edited."
        ),
        "bibliography": [
            "Smith J., 2023, Journal of Stuff, 10.1234/abc",
            "Aydın E. et al., 2024, IEEE TPAMI, 10.1109/xyz",
        ],
    }
    blob = module.export_docx(proposal)
    doc = Document(BytesIO(blob))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "1. Excellence" in headings
    assert "2. Impact" in headings
    assert "3. Implementation" in headings
    assert "Use of AI Tools in This Proposal" in headings
    assert "References" in headings

    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "GREENMOBILITY" in body_text
    assert "GMOB" in body_text
    assert "Of 412 sentences in the proposal" in body_text
    assert "Smith J., 2023" in body_text


def test_export_xlsx_budget_returns_none_for_now() -> None:
    """Lump Sum XLSX is the next task; v1 returns None."""

    module = get_module("horizon_eu_ria")
    assert module.export_xlsx_budget({"budget": {}}) is None


def test_template_builder_smoke() -> None:
    doc = build_template()
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Horizon Europe" in text
