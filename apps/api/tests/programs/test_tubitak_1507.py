"""TÜBİTAK 1507 (KOBİ AR-GE Başlangıç) module tests.

Plugin contract + 1507-specific blockers (KOBİ requirement, budget cap,
12-24 month duration window). Most of the validation engine is
inherited from :class:`TUBITAKBaseModule` and proven via the 1501
suite — these tests focus on the parameters that DIFFER from 1501.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from src.programs import REGISTRY, TUBITAK1507Module, get_module
from src.programs._tubitak_base import TUBITAKBaseModule
from src.programs.base import CallMetadata
from src.programs.tubitak_1507.template_builder import (
    build_template,
    to_bytes,
)

# ── Registry + identity ─────────────────────────────────────────────────


def test_module_in_registry() -> None:
    assert "tubitak_1507" in REGISTRY
    assert isinstance(get_module("tubitak_1507"), TUBITAK1507Module)


def test_inherits_from_tubitak_base() -> None:
    assert issubclass(TUBITAK1507Module, TUBITAKBaseModule)


def test_template_path_exists() -> None:
    module = get_module("tubitak_1507")
    assert Path(module.get_template_path()).is_file()


def test_prompt_paths_resolve_to_tubitak_1501_directory() -> None:
    """1507 reuses 1501's prompts — verify the path resolves there."""

    module = get_module("tubitak_1507")
    for agent_id in ("excellence_writer",):  # only 1501 prompt that exists today
        path = Path(module.get_prompt_path(agent_id))
        assert "tubitak_1501" in path.parts
        assert path.is_file(), f"missing prompt: {path}"


# ── Parameter overrides ────────────────────────────────────────────────


def test_class_parameters_diverge_from_1501() -> None:
    module = get_module("tubitak_1507")
    assert module.duration_min_months == 12
    assert module.duration_max_months == 24
    assert module.requires_kobi is True
    assert module.budget_max_tl == 1_500_000


# ── parse_call ─────────────────────────────────────────────────────────


def test_parse_call_normalises_to_turkish() -> None:
    module = get_module("tubitak_1507")
    metadata = module.parse_call("(test)", {"scope_summary": "x"})
    assert isinstance(metadata, CallMetadata)
    assert metadata.language == "tr"


# ── Brief schema ───────────────────────────────────────────────────────


def test_brief_schema_inherits_tubitak_shape() -> None:
    module = get_module("tubitak_1507")
    schema = module.get_brief_schema()
    assert len(schema.sections) >= 1
    field_keys = {f.key for s in schema.sections for f in s.fields}
    assert {"title", "problem_statement", "proposed_solution", "duration_months"} <= field_keys


# ── validate_draft ──────────────────────────────────────────────────────


def _b2_long_enough() -> str:
    """800-word B2 body so the b2_too_short rule doesn't fire by accident."""

    return "## B2 Yenilikçi Yönler\n\n" + " ".join(["kelime"] * 850) + "\n"


def test_validate_draft_clean_returns_no_blockers_with_kobi_brief() -> None:
    module = get_module("tubitak_1507")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _b2_long_enough(),
        "brief": {
            "duration_months": 18,
            "is_sme": True,
        },
        "budget": {"total_tl": 1_000_000},
    }
    issues = module.validate_draft(draft, metadata)
    blockers = [i for i in issues if i.severity == "blocker"]
    assert blockers == []


def test_validate_draft_blocks_when_not_kobi() -> None:
    module = get_module("tubitak_1507")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _b2_long_enough(),
        "brief": {"duration_months": 18, "is_sme": False},
    }
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "not_kobi" and i.severity == "blocker" for i in issues)


def test_validate_draft_does_not_flag_kobi_when_flag_missing() -> None:
    """is_sme missing → assume SME (TÜBİTAK is self-attestation)."""

    module = get_module("tubitak_1507")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _b2_long_enough(),
        "brief": {"duration_months": 18},
    }
    issues = module.validate_draft(draft, metadata)
    assert not any(i.code == "not_kobi" for i in issues)


def test_validate_draft_blocks_when_budget_exceeds_cap() -> None:
    module = get_module("tubitak_1507")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _b2_long_enough(),
        "brief": {"duration_months": 18, "is_sme": True},
        "budget": {"total_tl": 2_000_000},  # over 1.5M TL cap
    }
    issues = module.validate_draft(draft, metadata)
    assert any(i.code == "budget_exceeded" and i.severity == "blocker" for i in issues)


def test_validate_draft_blocks_when_duration_outside_1507_range() -> None:
    """30 months would pass 1501 (18-36) but fails 1507 (12-24)."""

    module = get_module("tubitak_1507")
    metadata = module.parse_call("", {})
    draft = {
        "excellence_md": _b2_long_enough(),
        "brief": {"duration_months": 30, "is_sme": True},
    }
    issues = module.validate_draft(draft, metadata)
    invalid = [i for i in issues if i.code == "invalid_duration"]
    assert len(invalid) == 1
    assert "TÜBİTAK 1507" in invalid[0].message_tr


# ── DOCX export ─────────────────────────────────────────────────────────


def test_export_docx_renders_with_1507_template() -> None:
    module = get_module("tubitak_1507")
    proposal = {
        "title": "AKILLI-FABRIKA",
        "draft": {
            "excellence_md": "## B1 Proje Konusu\n\nKOBİ için AI üretim hattı.\n",
            "impact_md": "## C1 Ekonomik Kazanım\n\nTürkiye sanayi 4.0.\n",
            "implementation_md": "## D1 İş Paketleri\n\nWP1, WP2, WP3.\n",
        },
        "budget": {"by_category": {"M1": [{"description": "Personel", "total": 800_000}]}},
    }
    blob = module.export_docx(proposal)
    doc = Document(BytesIO(blob))

    body = "\n".join(p.text for p in doc.paragraphs)
    assert "AKILLI-FABRIKA" in body
    assert "B. Bilimsel ve Teknolojik Detaylar" in body
    assert "C. Yaygın Etki ve Pazar" in body
    assert "D. Uygulama" in body


def test_export_xlsx_budget_returns_none_for_now() -> None:
    module = get_module("tubitak_1507")
    assert module.export_xlsx_budget({"budget": {}}) is None


# ── Template builder ────────────────────────────────────────────────────


def test_template_builder_smoke_returns_document_with_1507_heading() -> None:
    doc = build_template()
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "TÜBİTAK 1507" in text
    assert "TÜBİTAK 1501" not in text  # cover heading is 1507-specific


def test_template_builder_to_bytes_returns_valid_docx() -> None:
    blob = to_bytes()
    doc = Document(BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "B1" in text and "B2" in text and "B3" in text and "B4" in text
