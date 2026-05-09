"""Markdown → DOCX converter unit tests."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from src.exports.md_to_docx import render_markdown


def _render(md: str) -> Document:  # type: ignore[valid-type]
    doc = Document()
    render_markdown(doc, md)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def test_renders_three_heading_levels() -> None:
    doc = _render("# H1\n\n## H2\n\n### H3\n")
    headings = [(p.text, p.style.name) for p in doc.paragraphs]  # type: ignore[attr-defined]
    levels_seen = {style for _, style in headings if "Heading" in style}
    assert "Heading 1" in levels_seen
    assert "Heading 2" in levels_seen
    assert "Heading 3" in levels_seen


def test_paragraph_with_bold_and_italic() -> None:
    doc = _render("This is **bold** and *italic* text.")
    paragraphs = [p for p in doc.paragraphs if p.text]  # type: ignore[attr-defined]
    assert len(paragraphs) == 1
    runs = paragraphs[0].runs
    assert any(r.text == "bold" and r.bold for r in runs)
    assert any(r.text == "italic" and r.italic for r in runs)
    assembled = "".join(r.text for r in runs)
    assert assembled == "This is bold and italic text."


def test_bullet_list_consecutive_lines() -> None:
    doc = _render("- first\n- second\n* third")
    items = [
        p
        for p in doc.paragraphs  # type: ignore[attr-defined]
        if p.style.name == "List Bullet" and p.text
    ]
    assert [p.text for p in items] == ["first", "second", "third"]


def test_numbered_list_consecutive_lines() -> None:
    doc = _render("1. one\n2. two\n3. three")
    items = [
        p
        for p in doc.paragraphs  # type: ignore[attr-defined]
        if p.style.name == "List Number" and p.text
    ]
    assert [p.text for p in items] == ["one", "two", "three"]


def test_paragraphs_separated_by_blank_lines() -> None:
    doc = _render("First paragraph.\n\nSecond paragraph.\n\nThird.")
    paragraphs = [p.text for p in doc.paragraphs if p.text]  # type: ignore[attr-defined]
    assert paragraphs == ["First paragraph.", "Second paragraph.", "Third."]


def test_mixed_block_types_round_trip() -> None:
    md = (
        "## Section A\n\n"
        "Intro paragraph with **emphasis**.\n\n"
        "- bullet 1\n- bullet 2\n\n"
        "1. numbered 1\n2. numbered 2\n\n"
        "Closing paragraph.\n"
    )
    doc = _render(md)
    text_blocks = [(p.text, p.style.name) for p in doc.paragraphs if p.text]  # type: ignore[attr-defined]
    assert ("Section A", "Heading 2") in text_blocks
    assert ("bullet 1", "List Bullet") in text_blocks
    assert ("bullet 2", "List Bullet") in text_blocks
    assert ("numbered 1", "List Number") in text_blocks
    assert ("numbered 2", "List Number") in text_blocks
    assert ("Closing paragraph.", "Normal") in text_blocks
