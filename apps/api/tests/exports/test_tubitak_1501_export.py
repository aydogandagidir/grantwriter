"""TÜBİTAK 1501 export module tests.

Verifies that ``TUBITAK1501Module.export_docx`` produces a DOCX whose
content (re-opened with python-docx) contains every expected section,
the rendered Markdown body, and the budget table populated from
``proposal['budget']['by_category']``.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from src.programs import REGISTRY, TUBITAK1501Module, get_module
from src.programs.tubitak_1501.template_builder import build_template, to_bytes


def _sample_proposal() -> dict[str, object]:
    return {
        "id": "11111111-2222-3333-4444-555555555555",
        "tenant_id": "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee",
        "programme_id": "tubitak_1501",
        "title": "Tekstil Hattında Gerçek-Zamanlı Kalite Tespiti",
        "draft": {
            "excellence_md": (
                "## B1 Proje Konusu ve Amaçları\n\n"
                "Edge cihazlarda **gerçek-zamanlı** kalite tespiti.\n\n"
                "## B2 Yenilikçi Yönler\n\n"
                "Mevcut çözümler *yavaş*. Bizimki:\n\n"
                "- 200 ms gecikme\n- yerel veri işleme\n\n"
                "## B3 Yöntem ve Teknik\n\n"
                "1. Veri toplama\n2. Eğitim\n3. Saha pilotu\n\n"
                "## B4 Literatür Taraması\n\n"
                "Türkiye'de saha entegrasyonu az çalışılmıştır.\n"
            ),
            "impact_md": (
                "## C1 Ekonomik ve Ulusal Kazanım\n\n"
                "İhracat hacminde artış.\n\n"
                "## C2 Yaygın Etki\n\n"
                "Sektör genelinde dağılım.\n\n"
                "## C3 Pazar Analizi\n\n"
                "5000+ KOBİ adreslenebilir.\n"
            ),
            "implementation_md": (
                "## D1 İş Paketleri\n\nWP1-WP5 tanımları.\n\n"
                "## D2 Zaman Planlaması\n\n24 ay.\n\n"
                "## D3 Bütçe\n\nDetaylı bütçe tablosu altta.\n\n"
                "## D4 Proje Yönetimi ve Riskler\n\nRisk matrisi.\n"
            ),
        },
        "budget": {
            "by_category": {
                "M1": [
                    {
                        "description": "Kıdemli AR-GE Mühendisi",
                        "quantity": "12 adam-ay",
                        "unit_price": 35_000,
                        "total": 420_000,
                        "month": "1-12",
                    }
                ],
                "M2": [
                    {
                        "description": "Edge TPU geliştirme kiti",
                        "quantity": 4,
                        "unit_price": 2_500,
                        "total": 10_000,
                        "month": "1",
                    }
                ],
                "M3": [
                    {
                        "description": "Bulut altyapı",
                        "quantity": "12 ay",
                        "unit_price": 1_500,
                        "total": 18_000,
                        "month": "1-12",
                    }
                ],
            }
        },
    }


# ── Module + registry shape ─────────────────────────────────────────────


def test_module_in_registry() -> None:
    assert "tubitak_1501" in REGISTRY
    assert isinstance(get_module("tubitak_1501"), TUBITAK1501Module)


def test_template_path_exists() -> None:
    module = get_module("tubitak_1501")
    from pathlib import Path

    assert Path(module.get_template_path()).is_file()


def test_template_builder_produces_valid_docx() -> None:
    """Re-running the builder should produce a parseable docx with section
    placeholders. Catches regressions in the binary scaffold."""

    blob = to_bytes()
    doc = Document(BytesIO(blob))
    text = "\n".join(p.text for p in doc.paragraphs)  # type: ignore[attr-defined]
    assert "TÜBİTAK 1501" in text
    assert "B1 Proje Konusu" in text
    assert "C1 Ekonomik" in text
    assert "D1 İş Paketleri" in text


# ── export_docx end-to-end ──────────────────────────────────────────────


def test_export_docx_contains_all_sections_and_subsections() -> None:
    module = get_module("tubitak_1501")
    blob = module.export_docx(_sample_proposal())
    doc = Document(BytesIO(blob))

    headings = [
        p.text
        for p in doc.paragraphs  # type: ignore[attr-defined]
        if p.style.name.startswith("Heading")
    ]
    body_text = "\n".join(p.text for p in doc.paragraphs)  # type: ignore[attr-defined]

    # H1 section dividers, programmatically inserted.
    assert "B. Bilimsel ve Teknolojik Detaylar" in headings
    assert "C. Yaygın Etki ve Pazar" in headings
    assert "D. Uygulama" in headings
    assert "Bütçe Tablosu" in headings

    # H2 subsections, rendered from each section's markdown.
    for sub in (
        "B1 Proje Konusu ve Amaçları",
        "B2 Yenilikçi Yönler",
        "B3 Yöntem ve Teknik",
        "B4 Literatür Taraması",
        "C1 Ekonomik ve Ulusal Kazanım",
        "C2 Yaygın Etki",
        "C3 Pazar Analizi",
        "D1 İş Paketleri",
        "D2 Zaman Planlaması",
        "D3 Bütçe",
        "D4 Proje Yönetimi ve Riskler",
    ):
        assert sub in headings, f"missing heading: {sub}"

    # Proposal title carried through the template placeholder.
    assert "Tekstil Hattında Gerçek-Zamanlı Kalite Tespiti" in body_text


def test_export_docx_renders_inline_formatting_and_lists() -> None:
    module = get_module("tubitak_1501")
    blob = module.export_docx(_sample_proposal())
    doc = Document(BytesIO(blob))

    bullet_texts = [
        p.text
        for p in doc.paragraphs  # type: ignore[attr-defined]
        if p.style.name == "List Bullet" and p.text
    ]
    numbered_texts = [
        p.text
        for p in doc.paragraphs  # type: ignore[attr-defined]
        if p.style.name == "List Number" and p.text
    ]
    assert "200 ms gecikme" in bullet_texts
    assert "yerel veri işleme" in bullet_texts
    assert "Veri toplama" in numbered_texts

    # Bold + italic survived inline-formatting parsing.
    bold_runs = [
        run.text
        for p in doc.paragraphs  # type: ignore[attr-defined]
        for run in p.runs
        if run.bold
    ]
    italic_runs = [
        run.text
        for p in doc.paragraphs  # type: ignore[attr-defined]
        for run in p.runs
        if run.italic
    ]
    assert "gerçek-zamanlı" in bold_runs
    assert "yavaş" in italic_runs


def test_export_docx_renders_budget_table() -> None:
    module = get_module("tubitak_1501")
    blob = module.export_docx(_sample_proposal())
    doc = Document(BytesIO(blob))

    # Locate the budget table — the last table in the document.
    assert doc.tables, "no tables rendered"
    budget_table = doc.tables[-1]

    rows_text = [[cell.text for cell in row.cells] for row in budget_table.rows]
    header_row = rows_text[0]
    assert "Gider Kalemi" in header_row[0]
    assert "Toplam (TL)" in header_row[4]

    # Each populated category contributes one row — 3 in the sample.
    body_rows = rows_text[1:]
    assert any("M1" in row and "Kıdemli AR-GE Mühendisi" in row[1] for row in body_rows)
    assert any("M2" in row and "Edge TPU" in row[1] for row in body_rows)
    assert any("M3" in row and "Bulut altyapı" in row[1] for row in body_rows)

    # Last row is the grand-total summary; 420k + 10k + 18k = 448 000.
    total_row = rows_text[-1]
    assert total_row[0] == "TOPLAM"
    assert "448,000.00" in total_row[4] or "448.000,00" in total_row[4]


def test_export_docx_handles_missing_sections_gracefully() -> None:
    """A proposal where a writer hasn't run yet still exports — the
    missing section gets a placeholder line, not a crash."""

    proposal = _sample_proposal()
    proposal["draft"] = {"excellence_md": proposal["draft"]["excellence_md"]}  # type: ignore[index]
    proposal["budget"] = {"by_category": {}}

    blob = TUBITAK1501Module().export_docx(proposal)
    doc = Document(BytesIO(blob))
    body_text = "\n".join(p.text for p in doc.paragraphs)  # type: ignore[attr-defined]
    assert "[impact_md bölümü henüz üretilmedi.]" in body_text
    assert "[implementation_md bölümü henüz üretilmedi.]" in body_text


# ── Validation surface ──────────────────────────────────────────────────


def test_validate_draft_flags_short_b2() -> None:
    module = get_module("tubitak_1501")
    metadata = module.parse_call("(test fixture)", {})
    short_excellence = "## B2 Yenilikçi Yönler\n\nÇok kısa bir B2 metni.\n"
    issues = module.validate_draft({"excellence_md": short_excellence}, metadata)
    assert any(issue.code == "b2_too_short" and issue.severity == "blocker" for issue in issues)


def test_template_builder_smoke_returns_document() -> None:
    """Sanity-check the in-memory builder so a regression there doesn't
    silently leak into committed binary."""

    doc = build_template()
    text = "\n".join(p.text for p in doc.paragraphs)  # type: ignore[attr-defined]
    assert "TÜBİTAK 1501" in text
