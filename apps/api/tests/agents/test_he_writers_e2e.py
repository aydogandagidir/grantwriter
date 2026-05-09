"""End-to-end test for the HE writer chain.

Runs Excellence → Impact → Implementation in sequence with mocked LLM
responses, accumulates outputs into a draft dict, and feeds it to
``HorizonEURIAModule.export_docx``. Asserts the final DOCX has every
expected section.

The LLM is not contacted — :class:`tests.llm.conftest.FakeProvider`
returns canned Markdown for each writer task. This is the contract
test for the orchestrator's saga (S2.D7).
"""

from __future__ import annotations

import uuid
from io import BytesIO

from docx import Document
from src.agents import ExcellenceWriter, ImpactWriter, ImplementationWriter
from src.agents.base import AgentInput
from src.programs import get_module

from tests.llm.conftest import FakeProvider, build_router, make_response

# Canned Markdown bodies — minimal but cover all required subsections.
CANNED_EXCELLENCE = """\
## 1.1 Objectives and ambition

We will reach TRL 6 in 24 months, validated on three pilots. Gender
dimension is integrated through balanced data sampling and inclusive
team composition.

## 1.2 Methodology

A theory-of-change links data ingestion → model training → field
deployment, with explicit mitigations per stage.

## 1.3 State of the art

Prior work [Smith 2023] addressed the lab setting; we extend to
production [Aydın et al. 2024].

## 1.4 Open science practices

FAIR principles + DMP committed at month 3; code under EUPL.
"""

CANNED_IMPACT = """\
## 2.1 Project's pathways towards impact

KIPs: 25% reduction in defects, 200 SMEs adopting in 5 years, validated
via control-group sampling.

## 2.2 Measures to maximise impact

Dissemination via OPC UA standards body; exploitation through partner
B's industrial channel; communication via case-study videos. DNSH
considered for all six EU environmental objectives.

## 2.3 Summary canvas (key impact pathways)

- Specific needs: 1 200 SMEs in EU tooling sector
- Outcomes: validated runtime + standards-track contribution
- D&E&C measures: as detailed in §2.2
- KPIs: defect rate, adoption count, citations
"""

CANNED_IMPLEMENTATION = """\
## 3.1 Work plan and resources

WP1 (data) — Coord, 18 PM. WP2 (model) — Partner B, 22 PM.
WP3 (pilot) — Partner C, 16 PM. WP4 (dissemination) — all partners,
8 PM.

## 3.2 Capacity of the participants

Coordinator: 5 years AR-GE leadership. Partner B (Fraunhofer IPK):
synthetic-data pipelines for woven textiles.

## 3.3 Consortium as a whole

Three partners across NL / DE / TR balance industry + research.
Steering committee meets monthly; conflict resolution via independent
chair.

Risk register:
- Hardware delay (medium / high) — backup supplier identified.
- Training data shortage (low / medium) — synthetic-data fallback.
"""


def _make_router_with_three_responses() -> tuple[FakeProvider, FakeProvider]:
    primary = FakeProvider(
        "claude",
        [
            make_response(text=CANNED_EXCELLENCE, model="claude-opus-4-7", provider="claude"),
            make_response(text=CANNED_IMPACT, model="claude-opus-4-7", provider="claude"),
            make_response(text=CANNED_IMPLEMENTATION, model="claude-opus-4-7", provider="claude"),
        ],
    )
    fallback = FakeProvider("openai", [])
    return primary, fallback


async def test_he_three_writer_chain_renders_complete_part_b() -> None:
    primary, fallback = _make_router_with_three_responses()
    router = build_router(providers={"claude": primary, "openai": fallback})

    excellence = ExcellenceWriter(router=router)
    impact = ImpactWriter(router=router)
    implementation = ImplementationWriter(router=router)

    base_input = {
        "proposal_id": uuid.uuid4(),
        "tenant_id": uuid.uuid4(),
        "programme_id": "horizon_eu_ria",
        "language": "en",
        "brief": {
            "title": "GREENMOBILITY",
            "acronym": "GMOB",
            "problem_statement": "EU manufacturers depend on imported QC systems.",
            "proposed_solution": "Edge-deployed YOLOv9-tiny for on-line QC.",
            "trl_current": 4,
            "trl_target": 6,
            "duration_months": 24,
            "partners": "ACME (NL, SME)\nFraunhofer IPK (DE, RTO)\nGamma (TR, SME)\n",
        },
        "call": {"call_text": "HORIZON-CL4-2026 (test fixture)"},
    }

    call_analyst_output = {
        "agent_id": "call_analyst",
        "status": "completed",
        "output": {
            "scope_summary": "Industrial digital twins for manufacturing.",
            "key_terms_to_use": [
                "digital twin",
                "Edge AI",
                "FAIR",
                "TRL 6",
                "OPC UA",
            ],
            "page_limit": 45,
            "language_required": "en",
            "user_eligible": True,
        },
    }

    # Step 1: Excellence Writer.
    excellence_input = AgentInput(
        previous_outputs={"call_analyst": call_analyst_output, "rag_context": ""},
        **base_input,  # type: ignore[arg-type]
    )
    excellence_result = await excellence.run(excellence_input)
    assert excellence_result.status == "completed", excellence_result.metadata
    excellence_md = excellence_result.output["excellence_md"]
    assert "## 1.1" in excellence_md and "## 1.4" in excellence_md

    # Step 2: Impact Writer — feeds in the Excellence output.
    impact_input = AgentInput(
        previous_outputs={
            "call_analyst": call_analyst_output,
            "excellence_writer": {
                "agent_id": "excellence_writer",
                "status": "completed",
                "output": excellence_result.output,
            },
            "rag_context": "",
        },
        **base_input,  # type: ignore[arg-type]
    )
    impact_result = await impact.run(impact_input)
    assert impact_result.status == "completed", impact_result.metadata
    impact_md = impact_result.output["impact_md"]
    assert "## 2.1" in impact_md and "## 2.3" in impact_md

    # Step 3: Implementation Writer.
    implementation_input = AgentInput(
        previous_outputs={
            "call_analyst": call_analyst_output,
            "excellence_writer": {
                "agent_id": "excellence_writer",
                "status": "completed",
                "output": excellence_result.output,
            },
            "impact_writer": {
                "agent_id": "impact_writer",
                "status": "completed",
                "output": impact_result.output,
            },
            "rag_context": "",
        },
        **base_input,  # type: ignore[arg-type]
    )
    implementation_result = await implementation.run(implementation_input)
    assert implementation_result.status == "completed", implementation_result.metadata
    implementation_md = implementation_result.output["implementation_md"]
    for prefix in ("## 3.1", "## 3.2", "## 3.3"):
        assert prefix in implementation_md

    # Three LLM calls total — one per writer.
    assert len(primary.calls) == 3
    tasks_called = [req.task for req, _model, _key in primary.calls]
    assert tasks_called == ["excellence_writer", "impact_writer", "implementation_writer"]

    # Step 4: assemble draft + export DOCX.
    proposal = {
        "title": str(base_input["brief"]["title"]),
        "acronym": str(base_input["brief"]["acronym"]),
        "draft": {
            "excellence_md": excellence_md,
            "impact_md": impact_md,
            "implementation_md": implementation_md,
            "brief": base_input["brief"],
        },
    }
    module = get_module("horizon_eu_ria")
    blob = module.export_docx(proposal)
    doc = Document(BytesIO(blob))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "1. Excellence" in headings
    assert "2. Impact" in headings
    assert "3. Implementation" in headings
    for sub in (
        "1.1 Objectives and ambition",
        "1.4 Open science practices",
        "2.1 Project's pathways towards impact",
        "2.3 Summary canvas",  # truncated check; full label has more text
        "3.1 Work plan and resources",
        "3.3 Consortium as a whole",
    ):
        assert any(sub in h for h in headings), f"missing: {sub}"

    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "GREENMOBILITY" in body_text
    assert "GMOB" in body_text

    # validate_draft on the assembled draft must be clean (no blockers).
    metadata = module.parse_call("", call_analyst_output["output"])
    issues = module.validate_draft(proposal["draft"], metadata)
    blockers = [i for i in issues if i.severity == "blocker"]
    assert blockers == [], f"unexpected blockers: {[(b.code, b.message_en) for b in blockers]}"


async def test_impact_writer_falls_through_with_unknown_programme() -> None:
    primary = FakeProvider("claude", [])
    fallback = FakeProvider("openai", [])
    router = build_router(providers={"claude": primary, "openai": fallback})

    agent = ImpactWriter(router=router)
    result = await agent.run(
        AgentInput(
            proposal_id=uuid.uuid4(),
            tenant_id=uuid.uuid4(),
            programme_id="not_a_real_programme",  # truly unregistered
            language="en",
            brief={"title": "x"},
            call={"call_text": "x"},
            previous_outputs={},
        )
    )
    assert result.status == "failed"
    assert "unknown programme" in result.metadata["error"].lower()
    assert len(primary.calls) == 0
